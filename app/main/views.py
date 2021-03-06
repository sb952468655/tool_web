import os, sys, zipfile, time, logging, re, time, threading
from ..import create_app
from urllib.request import quote, unquote
from sqlite3 import OperationalError
from datetime import date, datetime
from flask import render_template, session, redirect, url_for, current_app, request, abort, g, Response, flash, app
from flask_login import login_required, current_user
import openpyxl
from openpyxl.styles import Alignment, PatternFill
import docx
from docx.shared import Pt
from docx.shared import RGBColor
from . import main
from ..check_pool import all_check
from ..inspection import mobile
from ..zuxun.zuxun import zuxun_check
from .config import g_city_to_name, g_log_path, g_backup_path, g_check_to_describe
from .gnenerate_config import parse_model, generate_config
from .address import get_address_data
from .statistic import get_statistic_data, get_statistic_host_data
from .. import db
from ..models import *
from .report import *
from .common import *
from ..special_line.views import save_special_line
from .forms import CaseUploadForm, ModelForm, ModelListCreateForm, ModelSelectForm, ConfigForm, PortServiceForm
sys.path.append('../')

@main.route('/')
@login_required
def index():
    session['action'] = 'report_port'
    if current_user.username == 'nokia' or current_user.username == 'sgs':
        return redirect(url_for('main.city_list'))
    else:
        session['city'] = current_user.username
        return redirect(url_for('main.report_port'))


@main.route('/report_port', methods=['GET', 'POST'])
@login_required
def report_port():
    '''统计报表-端口明细'''

    card_data = []
    page = request.args.get('page', 1, type=int)
    city = request.args.get('city')
    if not city:
        city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))
    else:
        session['city'] = city
        session['city_name'] = g_city_to_name.get(city)

    host_list = get_host(city)
    if not host_list:
        return '没有发现设备'
    search_date = date.today()
    form_date = ''
    
    from .report import get_report_data, get_device_name, get_port_ggl, get_ip

    if request.method == 'POST':
        host_name = request.form.get('host_name')
        form_date = request.form.get('date')
        if form_date:
            search_date = datetime.strptime(form_date,'%Y-%m-%d').date()
    else:
        host_name = request.args.get('host_name')
        if not host_name:
            host_name = host_list[0]

    if host_name == 'all':
        pageination = CardPort1.query.filter_by(city=city, date_time = search_date).paginate(
            page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
            error_out = False
        )
        if not pageination.items:
            last = CardPort1.query.filter_by(city=city).order_by(CardPort1.id.desc()).first()
            if last:
                pageination = CardPort1.query.filter_by(city=city, date_time = last.date_time).paginate(
                    page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
                    error_out = False
                )
    else:
        pageination = CardPort1.query.filter_by(host_name = host_name, date_time = search_date).paginate(
            page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
            error_out = False
        )

        if not pageination.items:
            last = CardPort1.query.filter_by(host_name = host_name).order_by(CardPort1.id.desc()).first()
            if last:
                pageination = CardPort1.query.filter_by(host_name = host_name, date_time = last.date_time).paginate(
                    page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
                    error_out = False
                )

        
    card_data = pageination.items

    card_data = [(index, item) for index, item in enumerate(card_data) ]
    return render_template('report.html', 
        card_data = card_data, 
        host_list = host_list,
        action = 'report_port', 
        host_name=host_name,
        date = form_date,
        pageination = pageination)


def save_host_list(city):
    '''设备清单入库'''

    # app = create_app('production')
    # app.app_context().push()
    host_list_count = HostList.query.filter_by(city = city, date_time = date.today()).count()

    if host_list_count:
        logging.info('host_list city: {} today is saved'.format(city))
        return

    host_list_data = get_host_list(city)

    logging.info('host_list city: {} begin save'.format(city))
    for item in host_list_data:
        host_list = HostList(
            city = city,
            host_name = item[0],
            host_ip = item[1],
            version = item[2],
            boot_time = item[3],
            config_save_time = item[4],
            date_time = date.today()
        )

        db.session.add(host_list)

    db.session.commit()
    db.session.close()

@main.route('/host_list_data', methods=['GET', 'POST'])
@login_required
def host_list_data():
    '''设备清单统计'''
    search_date = date.today()
    form_date = ''
    host_list_data = []
    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))

    if request.method == 'POST':
        form_date = request.form.get('date')
        if form_date:
            search_date = datetime.strptime(form_date,'%Y-%m-%d').date()

    host_list = get_host(city)
    for i in host_list:
        data = HostList.query.filter_by(host_name = i, date_time = search_date).first()
        if data:
            host_list_data.append(data)
        else:
            data = HostList.query.filter_by(host_name = i).order_by(HostList.id.desc()).first()
            if data:
                host_list_data.append(data)
    host_list_data = [(index, item) for index, item in enumerate(host_list_data) ]
    return render_template('host_list_data.html', host_list_data = host_list_data ,action='host_list_data')

@main.route('/card_detail', methods=['GET', 'POST'])
@login_required
def card_detail():
    '''card明细'''

    search_date = date.today()
    form_date = ''
    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))
    host_list = get_host(city)
    
    if request.method == 'POST':
        host_name = request.form.get('host_name')
        form_date = request.form.get('date')
        if form_date:
            search_date = datetime.strptime(form_date,'%Y-%m-%d').date()
    else:
        host_name = request.args.get('host_name')
        if not host_name:
            host_name = host_list[0]

    if host_name == 'all':
        card_detail_data = get_last_data(CardDetail, city=city, search_date=search_date)
    else:
        card_detail_data = get_last_data(CardDetail, host_name, search_date=search_date)

    card_detail_data = [(index, item) for index, item in enumerate(card_detail_data) ]
    return render_template('card_detail.html', host_name=host_name, host_list = host_list, card_detail_data = card_detail_data ,action='card_detail')

@main.route('/card_statistic', methods=['GET', 'POST'])
@login_required
def card_statistic():
    '''card统计'''

    search_date = date.today()
    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))
    host_list = get_host(city)

    if request.method == 'POST':
        host_name = request.form.get('host_name')
        form_date = request.form.get('date')
        if form_date:
            search_date = datetime.strptime(form_date,'%Y-%m-%d').date()
    else:
        host_name = request.args.get('host_name')
        if not host_name:
            host_name = host_list[0]

    if host_name == 'all':
        # card_statistic_data = CardStatistic.query.filter_by(city=city, date_time = search_date).all()
        # if not card_statistic_data:
        #     last = CardStatistic.query.filter_by(city=city).order_by(CardStatistic.id.desc()).first()
        #     if last:
        #         card_statistic_data = CardStatistic.query.filter_by(city=city, date_time = last.date_time).all()
        card_statistic_data = get_last_data(CardStatistic, city=city, search_date=search_date)
    else:
        # card_statistic_data = CardStatistic.query.filter_by(host_name = host_name, date_time = search_date).all()
        # if not card_statistic_data:
        #     last = CardStatistic.query.filter_by(host_name = host_name).order_by(CardStatistic.id.desc()).first()
        #     if last:
        #         card_statistic_data = CardStatistic.query.filter_by(host_name = host_name, date_time = last.date_time).all()
        card_statistic_data = get_last_data(CardStatistic, host_name, search_date=search_date)

    card_statistic_data = [(index, item) for index, item in enumerate(card_statistic_data) ]
    return render_template('card_statistic.html', host_name=host_name, host_list = host_list, card_statistic_data = card_statistic_data ,action='card_statistic')


@main.route('/all_card_statistic', methods=['GET', 'POST'])
@login_required
def all_card_statistic():
    '''全省板卡统计汇总'''

    card_statistic_data = []
    mda_statistic_data = []
    search_date = date.today()

    if request.method == 'POST':
        city = request.form.get('city')
        if not city:
            city = session.get('city')
        form_date = request.form.get('date')
        if form_date:
            search_date = datetime.strptime(form_date,'%Y-%m-%d').date()
    else:
        city = request.args.get('city')
        if not city:
            if current_user.username == 'nokia':
                city = 'all'
            else:
                city = session.get('city')

    if city == 'all':
        city_list = get_city_list()

        for i in city_list:
            card_statistic_data += get_last_data(CardStatistic, city=i, search_date=search_date)
            mda_statistic_data += get_last_data(MdaStatistic, city=i, search_date=search_date)

    else:
        card_statistic_data += get_last_data(CardStatistic, city=city, search_date=search_date)
        mda_statistic_data += get_last_data(MdaStatistic, city=city, search_date=search_date)
    
    data = [(index, item) for index, item in enumerate(card_statistic_data + mda_statistic_data) ]
    data_set = {}
    for _, i in data:
        if i.card_type not in data_set.keys():
            data_set[i.card_type] = int(i.card_num)
        else:
            data_set[i.card_type] = data_set[i.card_type] + int(i.card_num)

    data = []
    date_str = date.today().strftime('%Y-%m-%d')
    # if mda_statistic_data:
    #     date_str = mda_statistic_data[0].date_time
    for k, v in data_set.items():
        data.append((k,v))
    data = [(index, item) for index, item in enumerate(data) ]
    return render_template('all_card_statistic.html', 
    city = city, 
    date = date_str,
    city_list = g_city_to_name, 
    card_statistic_data = data, 
    action='all_card_statistic')


@main.route('/all_card_detail', methods=['GET', 'POST'])
@login_required
def all_card_detail():
    '''全省板卡统计明细'''
    
    card_data = []
    mda_data = []
    search_date = date.today()

    if request.method == 'POST':
        city = request.form.get('city')
        if not city:
            city = session.get('city')
        form_date = request.form.get('date')
        if form_date:
            search_date = datetime.strptime(form_date,'%Y-%m-%d').date()
    else:
        city = request.args.get('city')
        if not city:
            if current_user.username == 'nokia':
                city = 'all'
            else:
                city = session.get('city')

    if city == 'all':

        city_list = get_city_list()
        for i in city_list:
            card_data += get_last_data(CardDetail, city=i, search_date=search_date)
            mda_data += get_last_data(MdaDetail, city=i, search_date=search_date)

    else:
        card_data += get_last_data(CardDetail, city=city, search_date=search_date)
        mda_data += get_last_data(MdaDetail, city=city, search_date=search_date)
    
    for index, item in enumerate(mda_data):
        if item.slot == ' ':
            mda_data[index].slot = mda_data[index-1].slot.split('/')[0] + '/' + item.mda
        else:
            mda_data[index].slot = mda_data[index].slot.split('/')[0] + '/' + item.mda

    all_data = card_data + mda_data
    def take_host_name(elem):
        return elem.host_name
    all_data.sort(key=take_host_name)
    data = [(index, item) for index, item in enumerate(all_data) ]
    return render_template('all_card_detail.html', 
    city = city, 
    city_list = g_city_to_name, 
    card_detail_data = data, 
    action='all_card_detail')


@main.route('/mda_detail', methods=['GET', 'POST'])
@login_required
def mda_detail():
    '''mda明细'''

    search_date = date.today()
    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))
    host_list = get_host(city)

    if request.method == 'POST':
        host_name = request.form.get('host_name')
        form_date = request.form.get('date')
        if form_date:
            search_date = datetime.strptime(form_date,'%Y-%m-%d').date()
    else:
        host_name = request.args.get('host_name')
        if not host_name:
            host_name = host_list[0]

    if host_name == 'all':
        mda_detail_data = get_last_data(MdaDetail, city=city, search_date=search_date)
    else:
        mda_detail_data = get_last_data(MdaDetail, host_name, search_date=search_date)

    mda_detail_data = [(index, item) for index, item in enumerate(mda_detail_data) ]
    return render_template('mda_detail.html', host_name=host_name, host_list = host_list, mda_detail_data = mda_detail_data ,action='mda_detail')

@main.route('/mda_statistic', methods=['GET', 'POST'])
@login_required
def mda_statistic():
    '''mda统计'''

    search_date = date.today()
    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))
    host_list = get_host(city)

    if request.method == 'POST':
        host_name = request.form.get('host_name')
        form_date = request.form.get('date')
        if form_date:
            search_date = datetime.strptime(form_date,'%Y-%m-%d').date()
    else:
        host_name = request.args.get('host_name')
        if not host_name:
            host_name = host_list[0]

    if host_name == 'all':
        mda_statistic_data = get_last_data(MdaStatistic, city=city, search_date=search_date)
    else:
        mda_statistic_data = get_last_data(MdaStatistic, host_name, search_date=search_date)

    mda_statistic_data = [(index, item) for index, item in enumerate(mda_statistic_data) ]
    return render_template('mda_statistic.html', host_name=host_name, host_list = host_list, mda_statistic_data = mda_statistic_data, action='mda_statistic')

@main.route('/nat_port', methods=['GET', 'POST'])
@login_required
def nat_port():
    '''nat port 统计'''

    search_date = date.today()
    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))
    host_list = get_host(city)

    if request.method == 'POST':
        host_name = request.form.get('host_name')
        form_date = request.form.get('date')
        if form_date:
            search_date = datetime.strptime(form_date,'%Y-%m-%d').date()
    else:
        host_name = request.args.get('host_name')
        if not host_name:
            host_name = host_list[0]

    if host_name == 'all':
        nat_port_data = get_last_data(NatPort, city=city, search_date=search_date)
    else:
        nat_port_data = get_last_data(NatPort, host_name, search_date=search_date)

    nat_port_data = [(index, item) for index, item in enumerate(nat_port_data) ]
    return render_template('report/nat_port.html', host_name=host_name, host_list = host_list, data = nat_port_data, action='nat_port')

@main.route('/port_statistic', methods=['GET', 'POST'])
@login_required
def port_statistic():
    '''端口统计'''

    search_date = date.today()
    city = session.get('city')
    host_list = get_host(session.get('city'))
    if not city:
        return redirect(url_for('main.city_list'))

    if request.method == 'POST':
        host_name = request.form.get('host_name')
        form_date = request.form.get('date')
        if form_date:
            search_date = datetime.strptime(form_date,'%Y-%m-%d').date()
    else:
        host_name = request.args.get('host_name')
        if not host_name:
            host_name = host_list[0]

    if host_name == 'all':
        port_statistic_data = get_last_data(PortStatistic, city=city, search_date=search_date)
    else:
        port_statistic_data = get_last_data(PortStatistic, host_name, search_date=search_date)

    port_statistic_data = [(index, item) for index, item in enumerate(port_statistic_data) ]
    return render_template('port_statistic.html', host_name=host_name, host_list = host_list, port_statistic_data = port_statistic_data, action='port_statistic')


@main.route('/check_config/<host_name>')
@login_required
def check_config(host_name):
    '''配置检查'''

    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))

    config = get_log_first(city, host_name)
    if not config:
        abort(404)

    check_res = all_check.all_check(config)

    session['host_name'] = host_name
    return render_template('check.html', check_data=check_res, host_name=host_name)

@main.route('/upload_check', methods=['GET', 'POST'])
@login_required
def upload_check():
    '''配置上传检查'''

    if not os.path.exists(os.path.join('app', 'static', 'config')):
        os.makedirs(os.path.join('app', 'static', 'config'))
        
    form = ConfigForm()
    if form.validate_on_submit():
        f = form.upload_file.data
        save_path = os.path.join('app', 'static', 'config', f.filename)
        f.save(save_path)
        config = open(save_path).read()
        check_res = all_check.all_check(config)

        host_name = get_host_name(config)
        return render_template('check.html', check_data=check_res, host_name=host_name)

    return render_template('host_list_base.html', form=form)

@main.route('/check_excel/<host_name>')
@login_required
def check_excel(host_name):
    '''配置检查生成表格'''

    #先获取检查结果
    city = session.get('city')
    config = get_log(city, host_name)
    if not config:
        abort(404)
    check_res = all_check.all_check(config)
    
    file_name = '{}_配置检查'.format(host_name.split('.')[0])
    # excel.save(os.path.join('app','static', file_name))
    f = open(os.path.join('app','static', file_name), 'w')
    for item in check_res:
        f.write(item[0]+'\n\n'+item[1]+'\n\n')
    f.close()

    url = url_for('static', filename = file_name)
    url2 = unquote(url, encoding='utf-8')

    return redirect(url2)


@main.route('/xunjian/<host_name>')
@login_required
def xunjian(host_name):
    '''巡检'''

    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))

    xunjian_data = XunJian.query.filter_by(host_name = host_name, date_time = date.today()).all()
    if not xunjian_data:
        last = XunJian.query.filter_by(host_name = host_name).order_by(XunJian.id.desc()).first()
        if last:
            xunjian_data = XunJian.query.filter_by(host_name = host_name, date_time = last.date_time).all()
    warn_data = [item for item in xunjian_data if item.err]
    
    return render_template('xunjian/xunjian.html', xunjian_data=warn_data, host_name = host_name, city = city, check_to_describe = g_check_to_describe)

@main.route('/xunjian_output_all/<host_name>')
@login_required
def xunjian_output_all(host_name):
    '''设备巡检-全量输出'''

    city = session.get('city')
    xunjian_data = XunJian.query.filter_by(host_name = host_name, date_time = date.today()).all()
    if not xunjian_data:
        last = XunJian.query.filter_by(host_name = host_name).order_by(XunJian.id.desc()).first()
        if last:
            xunjian_data = XunJian.query.filter_by(host_name = host_name, date_time = last.date_time).all()

    warn_data = [item for item in xunjian_data if item.err]
    
    return render_template('xunjian/xunjian_output_all.html', xunjian_data=warn_data, host_name = host_name, city = city, check_to_describe = g_check_to_describe)

@main.route('/xunjian_all_host')
@login_required
def xunjian_all_host():
    '''所有设备巡检'''

    warn_data = []
    city = session.get('city')

    host_list = get_host(city)
    for i in host_list:
        xunjian_data = XunJian.query.filter_by(host_name = i, date_time = date.today()).all()
        if not xunjian_data:
            last = XunJian.query.order_by(XunJian.id.desc()).first()
            if last:
                xunjian_data = XunJian.query.filter_by(host_name = i, date_time = last.date_time).all()
        warn_data.append((i, [item for item in xunjian_data if item.err]))

    return render_template('xunjian/xunjian_all_host.html', xunjian_data=warn_data)

@main.route('/auto_config')
@login_required
def auto_config():
    '''脚本自动配置'''

    return render_template('auto_config.html')


@main.route('/generate_excel/<host_name>')
@login_required
def generate_excel(host_name):
    '''端口明细生成表格'''

    today_str = date.today().strftime('%Y%m%d')
    excel = openpyxl.Workbook()
    sheet = excel.active
    sheet['A1'] = '设备名'
    sheet['B1'] = '设备IP'
    sheet['C1'] = 'port'
    sheet['D1'] = '端口带宽'
    sheet['E1'] = 'Admin State'
    sheet['F1'] = 'Link State'
    sheet['G1'] = 'Port State'
    sheet['H1'] = 'CfgMTU'
    sheet['I1'] = 'OperMTU'
    sheet['J1'] = 'LAG'
    sheet['K1'] = 'PortMode'
    sheet['L1'] = 'PortEncp'
    sheet['M1'] = 'PortType'
    sheet['N1'] = 'C/QS/S/XFP/MDIMDX'
    sheet['O1'] = '收光功率'
    sheet['P1'] = '发光功率'
    sheet['Q1'] = '收光门限'
    sheet['R1'] = '发光门限'
    sheet['S1'] = '是否存在异常'

    sheet.column_dimensions['A'].width = 40.0
    sheet.column_dimensions['B'].width = 20.0
    sheet.column_dimensions['N'].width = 20.0
    sheet.column_dimensions['Q'].width = 15.0
    sheet.column_dimensions['R'].width = 15.0
    cur_row = 2
    count = CardPort1.query.filter_by(host_name = host_name, date_time = date.today()).count()
    if count == 0:
        abort(404)

    port_detail_data = CardPort1.query.filter_by(host_name = host_name, date_time = date.today()).all()
    for item in port_detail_data:
        sheet['A'+ str(cur_row)] = item.host_name
        sheet['B'+ str(cur_row)] = item.host_ip
        sheet['C'+ str(cur_row)] = item.port
        if '10' in item.port_dk:
            sheet['D'+ str(cur_row)] = '10GE'
        else:
            sheet['D'+ str(cur_row)] = 'GE'
        sheet['E'+ str(cur_row)] = item.admin_state
        sheet['F'+ str(cur_row)] = item.link_state
        sheet['G'+ str(cur_row)] = item.port_state
        sheet['H'+ str(cur_row)] = item.cfg_mtu
        sheet['I'+ str(cur_row)] = item.oper_mtu
        sheet['J'+ str(cur_row)] = item.lag
        sheet['K'+ str(cur_row)] = item.port_mode
        sheet['L'+ str(cur_row)] = item.port_encp
        sheet['M'+ str(cur_row)] = item.port_type
        sheet['N'+ str(cur_row)] = item.c_qs_s_xfp_mdimdx
        sheet['O'+ str(cur_row)] = item.optical_power

        if item.admin_state == 'Up' and item.link_state == 'Yes' and item.port_state == 'Up' and item.optical_power and (float(item.optical_power) < float(item.optical_warn.split('|')[1]) or float(item.optical_power) > float(item.optical_warn.split('|')[0])):
            fill = PatternFill("solid", fgColor="FF0000")
            sheet['O'+ str(cur_row)].fill = fill


        sheet['P'+ str(cur_row)] = item.output_power

        if item.admin_state == 'Up' and item.link_state == 'Yes' and item.port_state == 'Up' and item.output_power and (float(item.output_power) < float(item.output_warn.split('|')[1]) or float(item.output_power) > float(item.output_warn.split('|')[0])):
            fill = PatternFill("solid", fgColor="FF0000")
            sheet['P'+ str(cur_row)].fill = fill

        sheet['Q'+ str(cur_row)] = item.optical_warn
        sheet['R'+ str(cur_row)] = item.output_warn
        if item.admin_state == 'Up' and item.link_state == 'No' and item.port_state != 'Up':
            sheet['S'+ str(cur_row)] = '是'
            fill = PatternFill("solid", fgColor="FF0000")
            sheet['S'+ str(cur_row)].fill = fill
        else:
            sheet['S'+ str(cur_row)] = '否'

        cur_row += 1
    
    file_name = '端口明细-{}-{}.xlsx'.format(host_name, today_str)
    excel.save(os.path.join('app','static', file_name))

    return redirect(url_for('static', filename=file_name))

@main.route('/download_excel/<host_name>/<table_name>')
@login_required
def download_excel(host_name, table_name):
    '''生成excel，并生成下载链接'''

    data = []
    labels = []
    file_name = ''
    today_str = date.today().strftime('%Y%m%d')

    if table_name == 'port_statistic':
        file_name = '端口统计-{}-{}.xlsx'.format(host_name, today_str)
        labels = [
            '设备名', '设备IP', '端口类型', '数量', '已使用', '剩余'
        ]

        port_statistic_data = PortStatistic.query.filter_by(host_name = host_name, date_time = date.today()).all()
        for i in port_statistic_data:
            data.append((
                i.host_name,
                i.host_ip,
                i.port_type,
                i.port_num,
                i.used_num,
                i.unused_num
            ))

    elif table_name == 'card_detail':
        file_name = 'card明细-{}-{}.xlsx'.format(host_name, today_str)
        labels = [
            '设备名', '设备IP', 'Slot', 'Equipped Type', 'Admin State'
            , 'Operational State', 'Serial number', 'Time of last boot'
            , 'Temperature', 'Temperature threshold', '是否存在异常'
        ]

        card_detail_data = CardDetail.query.filter_by(host_name = host_name, date_time = date.today()).all()
        for i in card_detail_data:
            data.append((
                i.host_name,
                i.host_ip,
                i.slot,
                i.card_type,
                i.admin_state,
                i.operational_state,
                i.serial_number,
                i.time_of_last_boot,
                i.temperature,
                i.temperature_threshold,
                i.is_abnormal
            ))
    elif table_name == 'card_statistic':
        file_name = 'card统计-{}-{}.xlsx'.format(host_name, today_str)
        labels = [
            '设备名', '设备IP', 'card类型', '数量'
        ]
        card_statistic_data = []
        if host_name == 'all':
            city = session.get('city')
            host_list = get_host(city)
            for j in host_list:
                last = CardStatistic.query.filter_by(host_name = j).order_by(CardStatistic.id.desc()).first()
                if last:
                    temp_data = CardStatistic.query.filter_by(host_name = j, date_time = last.date_time).all()
                    card_statistic_data += temp_data
        else:
            last = CardStatistic.query.filter_by(host_name = host_name).order_by(CardStatistic.id.desc()).first()
            if last:
                card_statistic_data = CardStatistic.query.filter_by(host_name = host_name, date_time = last.date_time).all()
        # card_statistic_data = CardStatistic.query.filter_by(host_name = host_name, date_time = date.today()).all()
        for i in card_statistic_data:
            data.append((
                i.host_name,
                i.host_ip,
                i.card_type,
                i.card_num
            ))
    elif table_name == 'mda_detail':
        file_name = 'mda明细-{}-{}.xlsx'.format(host_name, today_str)
        labels = [
            '设备名', '设备IP', 'Slot', 'Mda', 'Equipped Type', 'Admin State'
            , 'Operational State', 'Serial number', 'Time of last boot'
            , 'Temperature', 'Temperature threshold', '是否存在异常'
        ]

        mda_detail_data = MdaDetail.query.filter_by(host_name = host_name, date_time = date.today()).all()
        for i in mda_detail_data:
            data.append((
                i.host_name,
                i.host_ip,
                i.slot,
                i.mda,
                i.equipped_type,
                i.admin_state,
                i.operational_state,
                i.serial_number,
                i.time_of_last_boot,
                i.temperature,
                i.temperature_threshold,
                i.is_abnormal
            ))
    elif table_name == 'mda_statistic':
        file_name = 'mda统计-{}-{}.xlsx'.format(host_name, today_str)
        labels = [
            '设备名', '设备IP', 'MDA类型', '数量'
        ]

        card_statistic_data = CardStatistic.query.filter_by(host_name = host_name, date_time = date.today()).all()
        for i in card_statistic_data:
            data.append((
                i.host_name,
                i.host_ip,
                i.card_type,
                i.card_num
            ))
    elif table_name == 'host_list':
        file_name = '设备清单统计-{}.xlsx'.format(today_str)
        city = session.get('city')

        labels = [
            '设备名', '设备IP', '版本', '设备启动时间', '设备配置保存时间'
        ]

        last = HostList.query.filter_by(city=city).order_by(HostList.id.desc()).first()
        if last:
            host_list_data = HostList.query.filter_by(city=city, date_time = last.date_time).all()

            for i in host_list_data:
                data.append((
                    i.host_name,
                    i.host_ip,
                    i.version,
                    i.boot_time,
                    i.config_save_time
                ))

    elif table_name == 'install_base':
        file_name = 'Install+base+统计表-{}-江苏.xlsx'.format(date.today().strftime('%Y%m%d'))
        labels = [
            '省份', '运营商', '业务类型', '网络类型', '设备类型', '设备型号', '版本', '数量', '备注'
        ]

        last = InstallBase.query.order_by(InstallBase.id.desc()).first()
        if last:
            install_base_data = InstallBase.query.filter_by(date_time = last.date_time).all()
            for i in install_base_data:
                data.append((
                    i.province,
                    i.operator,
                    i.busines_type,
                    i.net_type,
                    i.host_type,
                    i.host_model,
                    i.version,
                    i.number,
                    i.note
                ))
    elif table_name == 'net_flow':
        file_name = '重要网络流量统计-{} 江苏.xlsx'.format(date.today().strftime('%Y%m%d'))
        labels = [
            'Carrier', 'Province', 'City', 'Network', 'Site Name'
            , '10G Port Num', '10G Port Utilization', '100G Port Num'
            , '100G Port Utilization', 'Peak Uplink Throughput Utilization'
        ]

        last = NetFlow.query.order_by(NetFlow.id.desc()).first()
        if last:
            net_flow_data = NetFlow.query.filter_by(date_time = last.date_time).all()
            for i in net_flow_data:
                data.append((
                    i.carrier,
                    i.province,
                    i.city,
                    i.network,
                    i.site_name,
                    i.port_num_10g,
                    i.port_utilization_10g,
                    i.port_num_100g,
                    i.port_utilization_100g,
                    i.pea_uplink_throughput_utilization
                ))
    elif table_name == 'load_statistic':
        file_name = '业务负荷统计表（按端口）-{}-{}.xlsx'.format(host_name, today_str)
        labels = [
            '设备名', '设备IP', 'port', '端口带宽', '带宽in利用率'
            , '带宽out利用率', 'ies 3000用户数量', 'ies 3000地址池利用率'
            , 'vprn 4015用户数量', 'vprn 4015地址池利用率'
        ]

        city = session.get('city')
        load_statistic_data = LoadStatistic.query.filter_by(city = city, date_time = date.today()).all()
        for i in load_statistic_data:
            data.append((
                i.host_name,
                i.host_ip,
                i.port,
                i.port_dk,
                i.in_utilization,
                i.out_utilization,
                i.ies_3000_user_num,
                i.ies_3000_utilization,
                i.vprn_4015_user_num,
                i.vprn_4015_utilization
            ))
    elif table_name == 'load_statistic_host':
        file_name = '业务负荷统计表（按设备）-{}.xlsx'.format(today_str)
        labels = [
            '设备名', '设备IP', 'ies 3000用户数量', 'ies 3000地址池利用率', 'vprn 4015用户数量'
            , 'vprn 4015地址池利用率'
        ]

        city = session.get('city')
        load_statistic_host_data = LoadStatisticHost.query.filter_by(city = city, date_time = date.today()).all()
        for i in load_statistic_host_data:
            data.append((
                i.host_name,
                i.host_ip,
                i.ies_3000_num,
                i.ies_3000_pool_utilization,
                i.vprn_4015_num,
                i.vprn_4015_pool_utilization
            ))
    elif table_name == 'all_card_statistic':
        city = request.args.get('city')
        if not city:
            city = session.get('city')
        if city == 'all':
            file_name = '全省板卡统计汇总-{}.xlsx'.format(today_str)
        else:
            file_name = '{}板卡统计汇总-{}.xlsx'.format(g_city_to_name[city] ,today_str)
        labels = [
            '地市', '板卡类型', '板卡数量'
        ]


        if city == 'all':
            card_statistic_data = []
            mda_statistic_data = []
            city_list = get_city_list()
            for i in city_list:
                host_list = get_host(i)
                for j in host_list:
                    last = CardStatistic.query.filter_by(host_name = j).order_by(CardStatistic.id.desc()).first()
                    if last:
                        temp = CardStatistic.query.filter_by(host_name = j, date_time = last.date_time).all()
                        card_statistic_data += temp

                    last = MdaStatistic.query.filter_by(host_name = j).order_by(MdaStatistic.id.desc()).first()
                    if last:
                        temp = MdaStatistic.query.filter_by(host_name = j, date_time = last.date_time).all()
                        mda_statistic_data += temp
        else:
            # card_statistic_data = CardStatistic.query.filter_by(city = city, date_time = date.today()).all()
            # mda_statistic_data = MdaStatistic.query.filter_by(city = city, date_time = date.today()).all()

            card_statistic_data = []
            mda_statistic_data = []
            host_list = get_host(city)
            for j in host_list:
                
                last = CardStatistic.query.filter_by(host_name = j).order_by(CardStatistic.id.desc()).first()
                if last:
                    temp = CardStatistic.query.filter_by(host_name = j, date_time = last.date_time).all()
                    card_statistic_data += temp

                last = MdaStatistic.query.filter_by(host_name = j).order_by(MdaStatistic.id.desc()).first()
                if last:
                    temp = MdaStatistic.query.filter_by(host_name = j, date_time = last.date_time).all()
                    mda_statistic_data += temp


        temp = [(index, item) for index, item in enumerate(card_statistic_data + mda_statistic_data) ]
        data_set = {}
        for _, i in temp:
            if i.card_type not in data_set.keys():
                data_set[i.card_type] = int(i.card_num)
            else:
                data_set[i.card_type] = data_set[i.card_type] + int(i.card_num)

        res = []

        for k, v in data_set.items():
            res.append((k,v))

        for i in res:
            data.append((
                g_city_to_name[city],
                i[0],
                i[1]
            ))
    elif table_name == 'all_card_detail':

        city = request.args.get('city')
        if not city:
            city = session.get('city')
        if city == 'all':
            file_name = '全省板卡统计明细-{}.xlsx'.format(today_str)
        else:
            file_name = '{}板卡统计明细-{}.xlsx'.format(g_city_to_name[city] ,today_str)
        labels = [
            '地市', '设备名', '设备IP', 'slot', '板卡类型', 'admin状态', 'operational状态', 'serial number', 'time of last boot', 'temperature', 'temperature threshold'
        ]

        city = request.args.get('city')

        if city == 'all':
            card_data = []
            mda_data = []
            city_list = get_city_list()
            for i in city_list:
                host_list = get_host(i)
                for j in host_list:
                    last = CardDetail.query.filter_by(host_name = j).order_by(CardDetail.id.desc()).first()
                    if last:
                        temp = CardDetail.query.filter_by(host_name = j, date_time = last.date_time).all()
                        card_data += temp

                    last = MdaDetail.query.filter_by(host_name = j).order_by(MdaDetail.id.desc()).first()
                    if last:
                        temp = MdaDetail.query.filter_by(host_name = j, date_time = last.date_time).all()
                        mda_data += temp
        else:
            card_data = []
            mda_data = []
            host_list = get_host(city)
            for j in host_list:
                last = CardDetail.query.filter_by(host_name = j).order_by(CardDetail.id.desc()).first()
                if last:
                    temp = CardDetail.query.filter_by(host_name = j, date_time = last.date_time).all()
                    card_data += temp

                last = MdaDetail.query.filter_by(host_name = j).order_by(MdaDetail.id.desc()).first()
                if last:
                    temp = MdaDetail.query.filter_by(host_name = j, date_time = last.date_time).all()
                    mda_data += temp

        for index, item in enumerate(mda_data):
            if item.slot == ' ':
                mda_data[index].slot = mda_data[index-1].slot.split('/')[0] + '/' + item.mda
            else:
                mda_data[index].slot = mda_data[index].slot.split('/')[0] + '/' + item.mda

        for i in card_data:
            data.append((
                g_city_to_name[i.city],
                i.host_name,
                i.host_ip,
                i.slot,
                i.card_type,
                i.admin_state,
                i.operational_state,
                i.serial_number,
                i.time_of_last_boot,
                i.temperature,
                i.temperature_threshold
            ))

        for i in mda_data:
            data.append((
                g_city_to_name[i.city],
                i.host_name,
                i.host_ip,
                i.slot,
                i.equipped_type,
                i.admin_state,
                i.operational_state,
                i.serial_number,
                i.time_of_last_boot,
                i.temperature,
                i.temperature_threshold
            ))
    elif table_name == 'nat_port':
        file_name = 'Nat Port 利用率统计表-{}.xlsx'.format(today_str)
        labels = [
            '设备名', 'Port', '利用率', '采集日期'
        ]

        nat_port_data = NatPort.query.filter_by(host_name = host_name, date_time = date.today()).all()
        for i in nat_port_data:
            data.append((
                i.host_name,
                i.port,
                i.utilization,
                i.date_time
            ))

    save_path = os.path.join('app','static', file_name)
    if labels and data:
        make_excel(save_path, labels, data)
        return redirect(url_for('static', filename=file_name))
    else:
        abort(404)

@main.route('/xj_log_export/<host_name>')
@login_required
def xj_log_export(host_name):
    '''巡检日志导出'''

    city = session.get('city')
    config = get_log(city, host_name)
    if not config:
        abort(404)

    log_config_str = config[config.index('# Finished'):]
    with open(os.path.join(g_backup_path, host_name + '.log'), 'w') as f:
        f.write(log_config_str)

    return redirect(url_for('static', filename = 'backup/{}.log'.format(host_name)))

@main.route('/xj_log/<host_name>')
@login_required
def xj_log(host_name):
    '''巡检log全量输出'''

    city = session.get('city')
    config = get_log(city, host_name)
    if not config:
        abort(404)

    log_str = config[config.index('# Finished'):]

    return render_template('xunjian/xunjian_log.html', log_str = log_str, host_name = host_name)

@main.route('/xj_report/<city>/<host_name>/<type>')
@login_required
def xj_report(city, host_name, type):
    '''生成巡检报告'''

    config = get_log(city, host_name)
    if not config:
        abort(404)

    # xunjian_data = mobile.xunjian(config, config)
    count = XunJian.query.filter_by(host_name = host_name, date_time = date.today()).count()
    if count == 0:
        abort(404)

    xunjian_data = XunJian.query.filter_by(host_name = host_name, date_time = date.today()).all()
    # warn_data = [item for item in xunjian_data if item.err]

    doc = docx.Document(os.path.join('app','static','xunjian.docx'))
    area = g_city_to_name.get(city) + '移动'
    doc.add_paragraph(area + '巡检报告', style='report-head')
    if type == '1':
        doc.add_paragraph('上海贝尔7750单设备巡检报告', style='report-head')
    else:
        doc.add_paragraph('上海贝尔7750单设备巡检输出', style='report-head')
    doc.add_paragraph()
    doc.add_paragraph()
    today_obj = datetime.now()
    today = '%d年%d月%d日' % (today_obj.year, today_obj.month, today_obj.day)
    doc.add_paragraph(today, style='report-date')
    doc.add_page_break()
    doc.add_heading('巡检情况汇总', 4)
    
    if type == '1':
        warn_data = [item for item in xunjian_data if item.err and '正常' not in item.err]
    else:
        warn_data = [item for item in xunjian_data if item.err]
    
    p_name = doc.add_paragraph(host_name.split('.')[0], style='report-info')
    font = p_name.runs[0].font
    font.color.rgb = RGBColor(0, 0, 0)
    for line in warn_data:
        p_info = doc.add_paragraph(line.check_item, style='report-info')
        p_warn = doc.add_paragraph('注：' + line.err, style='report-normal')

        doc.add_paragraph()
        
        font = p_info.runs[0].font
        if '正常' in line.err:
            font.color.rgb = RGBColor(0, 0, 255)
        else:
            font.color.rgb = RGBColor(244, 208, 0)
        # font.color.rgb = RGBColor(0, 0, 255)

        # font = p_warn.runs[0].font
        # font.color.rgb = RGBColor(255, 0, 255)


    doc.add_heading('总结', 4)

    doc.add_paragraph('1，为了保障%s移动城域网7750设备正常运行，请定期清理过滤网。' % g_city_to_name.get(city),
        style='report-normal')

    for item in warn_data:
        if '板卡温度高' in item.err:
            doc.add_paragraph('2，板卡温度高建议清洗防尘网。',
            style='report-normal')
        break

    if type == '1':
        report_name = '%s移动%s单设备巡检报告-%s.docx' % (g_city_to_name.get(city), host_name, today)
    else:
        report_name = '%s移动%s单设备巡检输出-%s.docx' % (g_city_to_name.get(city), host_name, today)

    doc.save( os.path.join('app', 'static', report_name))
    return redirect(url_for('static', filename = report_name))

@main.route('/xj_report_all_host')
@login_required
def xj_report_all_host(city=None):
    '''所有设备巡检，生成巡检报告'''


    warn_data = []
    if not city:
        city = session.get('city')

    today_obj = datetime.now()
    today = '%d年%d月%d日' % (today_obj.year, today_obj.month, today_obj.day)
    report_name = '%s移动巡检报告-%s.docx' % (g_city_to_name.get(city), today)
    file_path = os.path.join('app', 'static', report_name)
    if os.path.exists(file_path):
        return redirect(url_for('static', filename = report_name))
    host_list = get_host(city)
    for i in host_list:
        xunjian_data = XunJian.query.filter_by(host_name = i, date_time = date.today()).all()
        if not xunjian_data:
            last = XunJian.query.filter_by(host_name = i).order_by(XunJian.id.desc()).first()
            if last:
                xunjian_data = XunJian.query.filter_by(host_name = i, date_time = last.date_time).all()
        if xunjian_data:
            warn_data.append((i, [item for item in xunjian_data if item.err]))


    #生成报告
    doc = docx.Document(os.path.join('app','static','xunjian.docx'))

    area = g_city_to_name.get(city) + '移动'
    doc.add_paragraph(area + '巡检报告', style='report-head')
    doc.add_paragraph('上海贝尔7750设备巡检报告', style='report-head')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(today, style='report-date')
    doc.add_page_break()
    doc.add_heading('巡检情况汇总', 4)
    
    p_num = 1
    for line in warn_data:
        p_name = doc.add_paragraph(str(p_num) + '. ' + line[0] + '\n', style='report-info')
        p_num += 1
        font = p_name.runs[0].font
        font.color.rgb = RGBColor(0, 0, 0)

        for j in line[1]:
            
            p_info = doc.add_paragraph(j.check_item, style='report-info')
            p_warn = doc.add_paragraph('注：' + j.err, style='report-normal')

            doc.add_paragraph()
            
            font = p_info.runs[0].font
            if '正常' in j.err:
                font.color.rgb = RGBColor(0, 0, 255)
            else:
                font.color.rgb = RGBColor(244, 208, 0)

            font = p_warn.runs[0].font
            if 'port状态巡检' == j.check_item:
                font.size = Pt(9)
            else:
                font.size = Pt(12)
            # font.color.rgb = RGBColor(255, 0, 255)


    doc.add_heading('总结', 4)

    doc.add_paragraph('1，为了保障%s移动城域网7750设备正常运行，请定期清理过滤网。' % g_city_to_name.get(city),
        style='report-normal')

    for item in warn_data:
        for j in item[1]:
            if 'Temperature' in j.err:
                doc.add_paragraph('2，板卡温度高建议清洗防尘网。',
                style='report-normal')
            break
    
    
    doc.save( os.path.join('app', 'static', report_name))

    return redirect(url_for('static', filename = report_name))

@main.route('/xj_report_summary')
@login_required
def xj_report_summary():
    '''所有设备巡检报告汇总'''

    logging.info('所有设备巡检报告汇总 开始')
    all_data = []
    city_list = get_city_list()
    all_city_report_name = '全省巡检报告汇总-{}.xlsx'.format(date.today().strftime('%Y-%m-%d'))

    for i in city_list:
        city_data = []
        host_list = get_host(i)
        logging.info('city {} host_list {}'.format(i, host_list))
        for j in host_list:
            last = XunJian.query.filter_by(host_name = j).order_by(XunJian.id.desc()).first()
            if last:
                logging.info('host: {} last data ok'.format(j))
                data = XunJian.query.filter_by(host_name = j, date_time = last.date_time).all()
                # xunjian_data += data
                city_data += data
            else:
                logging.info('host: {} last data err')

    # file_name = '报告汇总.xlsx'
        labels = [
            '设备名', '检查项', '检查结果', '检查日期'
        ]

        data = []
        for k in city_data:
            check_res = '正常'
            if '正常' not in k.err:
                check_res = '异常'

            data.append((
                k.host_name,
                k.check_item,
                check_res,
                k.date_time.strftime('%Y-%m-%d')
            ))
            all_data.append((
                k.host_name,
                k.check_item,
                check_res,
                k.date_time.strftime('%Y-%m-%d')
            ))

        file_name = '巡检报告汇总-{}-{}.xlsx'.format(g_city_to_name[i], date.today().strftime('%Y-%m-%d'))
        save_path = os.path.join('app','static', file_name)
        if labels and data:
            make_excel(save_path, labels, data)
            # return redirect(url_for('static', filename=file_name))
        # else:
        #     abort(404)
    #生成全省报告
    save_path = os.path.join('app','static', all_city_report_name)
    if all_data:
        logging.info('xj_report_summary begin save')
        make_excel(save_path, labels, all_data)
    else:
        logging.info('xj_report_summary today not data')

@main.route('/xj_summary_download')
@login_required
def xj_summary_download():
    '''下载巡检报告汇总'''

    is_all = request.args.get('city')
    city = session.get('city')
    url = ''
    if is_all == 'all':
        file_name = '全省巡检报告汇总-{}.xlsx'.format(date.today().strftime('%Y-%m-%d'))
        file_path = os.path.join('app', 'static', file_name)
        if os.path.exists(file_path):
            url = url_for('static', filename = file_name)
    else:
        file_name = '巡检报告汇总-{}-{}.xlsx'.format(g_city_to_name[city], date.today().strftime('%Y-%m-%d'))
        file_path = os.path.join('app', 'static', file_name)
        if os.path.exists(file_path):
            url = url_for('static', filename = file_name)

    if url:
        return redirect(url)
    else:
        abort(404)

@main.route('/city_list')
@login_required
def city_list():
    '''地市列表'''

    session['city'] = ''
    city_data = []
    citys = get_city_list()
    for i in citys:
        city_data.append((i, g_city_to_name.get(i)))
    return render_template('city_list.html', city_data = city_data)


@main.route('/host_list/<action>')
@login_required
def host_list(action):
    session['action'] = action
    if not session.get('city'):
        session['city'] = request.args.get('city')
        session['city_name'] = session.get('city')

    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))

    session['city_name'] = g_city_to_name.get(city)
    
    form = ConfigForm()
    host_data = get_host(city)
    host_data = [(index, host) for index, host in enumerate(host_data)]
    if action == 'config_backup':
        return redirect(url_for('main.config_backup'))

    return render_template('host_list_base.html', host_data = host_data,  action = action, form=form)


@main.route('/address_collect', methods=['GET', 'POST'])
@login_required
def address_collect():
    '''三层接口和静态用户IP地址采集'''

    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))
    host_list = get_host(city)

    host_data = [(index, host) for index, host in enumerate(host_list)]
    return render_template('report/address_collect_download.html',
    host_data = host_data, action='address_collect')


@main.route('/address_mk_excel', methods=['POST'])
@login_required
def address_mk_excel():
    '''地址采集生成 excel 表格'''

    if not os.path.exists(os.path.join('app', 'static', 'address_collect')):
            os.makedirs(os.path.join('app', 'static', 'address_collect'))

    backup_data = request.form.to_dict()
    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))
    host_num = len(get_host(city))
    if host_num == len(backup_data):
        file_name = '地址采集-{}-{}.zip'.format(g_city_to_name.get(city), date.today().strftime('%Y%m%d'))
    elif len(backup_data) == 1:
        keys = list(backup_data.keys())
        file_name = '地址采集-{}-{}.zip'.format(keys[0], date.today().strftime('%Y%m%d'))
    else:
        file_name = '地址采集-' + str(time.time()) + '.zip'
    zip_name = os.path.join('app', 'static', 'address_collect', file_name)
    zip = zipfile.ZipFile(zip_name, 'w', zipfile.ZIP_DEFLATED )
    #获取选中的设备名称
    for host_name, _ in request.form.to_dict().items():
        save_path = os.path.join('app','static', 'address_collect', '{}-{}-地址采集.xlsx'.format(host_name, date.today().strftime('%Y%m%d')))
        if not os.path.exists(save_path):
            excel = openpyxl.Workbook()
            sheet = excel.active
            sheet['A1'] = '设备名'
            sheet['B1'] = '设备IP'
            sheet['C1'] = 'IP类型'
            sheet['D1'] = '网络功能类型'
            sheet['E1'] = '是否已分配'
            sheet['F1'] = 'IP'
            sheet['G1'] = '网关'
            sheet['H1'] = '掩码'
            sheet['I1'] = '逻辑接口编号'
            sheet['J1'] = 'sap-ID'
            sheet['K1'] = '下一跳IP'
            sheet['L1'] = 'IES/VPRN编号'
            sheet['M1'] = 'VPN-RD'
            sheet['N1'] = 'VPN-RT'
            sheet['O1'] = '接口或用户描述'


            sheet.column_dimensions['A'].width = 40.0
            sheet.column_dimensions['B'].width = 20.0
            # sheet.column_dimensions['N'].width = 20.0
            # sheet.column_dimensions['Q'].width = 15.0
            # sheet.column_dimensions['R'].width = 15.0
            cur_row = 2

            address_data = AddressCollect.query.filter_by(host_name = host_name, date_time = date.today()).all()
            for item in address_data:
                sheet['A'+ str(cur_row)] = item.host_name.split('.')[0]
                sheet['B'+ str(cur_row)] = item.host_ip
                sheet['C'+ str(cur_row)] = item.ip_type
                sheet['D'+ str(cur_row)] = item.function_type
                sheet['E'+ str(cur_row)] = item.is_use
                sheet['F'+ str(cur_row)] = item.ip
                sheet['G'+ str(cur_row)] = item.gateway
                sheet['H'+ str(cur_row)] = item.mask
                sheet['I'+ str(cur_row)] = item.interface_name
                sheet['J'+ str(cur_row)] = item.sap_id
                sheet['K'+ str(cur_row)] = item.next_hop
                sheet['L'+ str(cur_row)] = item.ies_vprn_id
                sheet['M'+ str(cur_row)] = item.vpn_rd
                sheet['N'+ str(cur_row)] = item.vpn_rt
                sheet['O'+ str(cur_row)] = item.description

                cur_row += 1
            
            excel.save(save_path)
        zip.write(save_path, '{}-{}-地址采集.xlsx'.format(host_name, date.today().strftime('%Y%m%d')))

    zip.close()
    url = url_for('static', filename = 'address_collect/{}'.format(file_name))
    url2 = unquote(url, encoding='utf-8')
    return redirect(url2)

@main.route('/service_mk_excel', methods=['POST'])
@login_required
def service_mk_excel():
    '''端口业务类型统计生成 excel 表格'''

    if not os.path.exists(os.path.join('app', 'static', 'service_statistic')):
        os.makedirs(os.path.join('app', 'static', 'service_statistic'))

    backup_data = request.form.to_dict()
    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))
    host_num = len(get_host(city))
    if host_num == len(backup_data):
        file_name = '端口业务统计-{}-{}.zip'.format(g_city_to_name.get(city), date.today().strftime('%Y%m%d'))
    elif len(backup_data) == 1:
        keys = list(backup_data.keys())
        file_name = '端口业务统计-{}-{}.zip'.format(keys[0], date.today().strftime('%Y%m%d'))
    else:
        file_name = '端口业务统计-' + str(time.time()) + '.zip'
    zip_name = os.path.join('app', 'static', 'service_statistic', file_name)
    zip = zipfile.ZipFile(zip_name, 'w', zipfile.ZIP_DEFLATED )
    #获取选中的设备名称
    for host_name, _ in request.form.to_dict().items():
        save_path = os.path.join('app','static', 'service_statistic', '{}-{}-端口业务统计.xlsx'.format(host_name, date.today().strftime('%Y%m%d')))
        if not os.path.exists(save_path):
            excel = openpyxl.Workbook()
            sheet = excel.active
            sheet['A1'] = '设备名'
            sheet['B1'] = '端口'
            sheet['C1'] = '端口描述'
            sheet['D1'] = 'lag-id'
            sheet['E1'] = 'sap/port'
            sheet['F1'] = 'interface'
            sheet['G1'] = 'service'
            sheet['H1'] = 'service id'
            sheet['I1'] = '采集日期'

            sheet.column_dimensions['A'].width = 40.0
            sheet.column_dimensions['E'].width = 15.0
            sheet.column_dimensions['I'].width = 15.0
            cur_row = 2

            service_statistic_data = []
            # last = ServiceStatistic.query.filter_by(host_name=host_name).order_by(ServiceStatistic.id.desc()).first()
            # if last:
            #     service_statistic_data = ServiceStatistic.query.filter_by(host_name=host_name, date_time = last.date_time).all()
            config = get_log_first(city, host_name)
            if not config:
                continue
            service_statistic_data = get_service_statistic(config)
            for item in service_statistic_data:
                sheet['A'+ str(cur_row)] = host_name
                sheet['B'+ str(cur_row)] = item[0]
                sheet['C'+ str(cur_row)] = item[1]
                sheet['D'+ str(cur_row)] = item[2]
                sheet['E'+ str(cur_row)] = item[3]
                sheet['F'+ str(cur_row)] = item[4]
                sheet['G'+ str(cur_row)] = item[5]
                sheet['H'+ str(cur_row)] = item[6]
                sheet['I'+ str(cur_row)] = datetime.today().strftime('%Y-%m-%d')

                cur_row += 1
            
            excel.save(save_path)
        zip.write(save_path, '{}-{}-端口业务统计.xlsx'.format(host_name, date.today().strftime('%Y-%m-%d')))

    zip.close()
    url = url_for('static', filename = 'service_statistic/{}'.format(file_name))
    url2 = unquote(url, encoding='utf-8')
    return redirect(url2)

@main.route('/upload_service_statistic', methods=['POST'])
@login_required
def upload_service_statistic():
    '''根据上传文件统计端口业务类型 生成 excel 表格'''

    if not os.path.exists(os.path.join('app', 'static', 'config')):
        os.makedirs(os.path.join('app', 'static', 'config'))
        
    form = PortServiceForm()
    if form.validate_on_submit():
        f = form.upload_file.data
        save_path = os.path.join('app', 'static', 'config', f.filename)
        f.save(save_path)
        try:
            config = open(save_path).read()
            host_name = get_host_name(config)
        except:
            flash('配置异常，请检查！')
            return redirect(url_for('main.service_statistic'))
        file_name = '{}-{}-端口业务统计.xlsx'.format(host_name, date.today().strftime('%Y%m%d'))
        save_path = os.path.join('app','static', 'service_statistic', file_name)
        excel = openpyxl.Workbook()
        sheet = excel.active
        sheet['A1'] = '设备名'
        sheet['B1'] = '端口'
        sheet['C1'] = '端口描述'
        sheet['D1'] = 'lag-id'
        sheet['E1'] = 'sap/port'
        sheet['F1'] = 'interface'
        sheet['G1'] = 'service'
        sheet['H1'] = 'service id'
        sheet['I1'] = '采集日期'

        sheet.column_dimensions['A'].width = 40.0
        sheet.column_dimensions['E'].width = 15.0
        sheet.column_dimensions['I'].width = 15.0
        cur_row = 2

        service_statistic_data = []
        service_statistic_data = get_service_statistic(config)
        for item in service_statistic_data:
            sheet['A'+ str(cur_row)] = host_name
            sheet['B'+ str(cur_row)] = item[0]
            sheet['C'+ str(cur_row)] = item[1]
            sheet['D'+ str(cur_row)] = item[2]
            sheet['E'+ str(cur_row)] = item[3]
            sheet['F'+ str(cur_row)] = item[4]
            sheet['G'+ str(cur_row)] = item[5]
            sheet['H'+ str(cur_row)] = item[6]
            sheet['I'+ str(cur_row)] = datetime.today().strftime('%Y-%m-%d')

            cur_row += 1
        
        excel.save(save_path)

        url = url_for('static', filename = 'service_statistic/{}'.format(file_name))
        url2 = unquote(url, encoding='utf-8')
        return redirect(url2)
    else:
        return redirect(url_for('main.service_statistic'))

@main.route('/config_backup')
@login_required
def config_backup():
    '''config备份（下载到本地）'''
    host_data = []
    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))

    host_list = get_host(city)
    # today = date.today()
    date_str = date.today().strftime('%Y%m%d')
    for item in host_list:
        # log1 = get_log(city, item)
        log1 = get_log_first(city, item)
        if log1:
            log_str = log1

        if log_str:
            date_str = get_log_date(log_str)
            host_data.append((item, date_str))
    host_data = [(index, host) for index, host in enumerate(host_data)]
    return render_template('back_up/config_backup_host_list.html', host_data=host_data)

@main.route('/backup_list', methods=['GET', 'POST'])
@login_required
def backup_list():
    '''获取需要备份的config列表'''

    city = session.get('city')
    select_host = []
    if not os.path.exists(os.path.join('app', 'static', 'backup')):
        os.makedirs(os.path.join('app', 'static', 'backup'))
    if request.method=='POST':
        host_num = len(get_host(city))
        backup_data = request.form.to_dict()
        if host_num == len(backup_data):
            file_name = 'config备份-{}-{}.zip'.format(g_city_to_name.get(city), date.today().strftime('%Y%m%d'))
        elif len(backup_data) == 1:
            keys = list(backup_data.keys())
            file_name = 'config备份-{}-{}.zip'.format(keys[0], date.today().strftime('%Y%m%d'))
        else:
            file_name = 'config备份-' + str(time.time()) + '.zip'
        select_host = [ (city, name) for name, _ in request.form.to_dict().items()]
    else:
        file_name = '全省config备份-{}.zip'.format(date.today().strftime('%Y%m%d'))

        city_list = get_city_list()
        for i in city_list:
            host = get_host(i)
            for j in host:
                select_host.append((i, j))

    zip_name = os.path.join('app', 'static', 'backup', file_name)
    zip = zipfile.ZipFile(zip_name, 'w', zipfile.ZIP_DEFLATED )
    for city, name in select_host:
        log_str = get_log_first(city, name)
        if log_str:
            log_config_str = log_str[:log_str.index('# Finished')]
            back_up_log_path = os.path.join('app', 'static', 'backup', 'config备份-{}-{}.log'.format(name, date.today().strftime('%Y%m%d')))
            with open(back_up_log_path, 'w') as f:
                f.write(log_config_str)
            zip.write(back_up_log_path, 'config备份-{}-{}.log'.format(name, date.today().strftime('%Y%m%d')))

    zip.close()

    url = url_for('static', filename = 'backup/{}'.format(file_name))
    url2 = unquote(url, encoding='utf-8')
    return redirect(url2)



@main.route('/load_statistic', methods=['POST', 'GET'])
@login_required
def load_statistic():
    '''业务负荷统计'''
    statistic_data = None
    page = request.args.get('page', 1, type=int)
    city = session.get('city')
    search_date = date.today()
    form_date = ''
    host_list = get_host(session.get('city'))
    if not city:
        return redirect(url_for('main.city_list'))

    if request.method == 'POST':
        host_name = request.form.get('host_name')
        form_date = request.form.get('date')
    else:
        host_name = host_list[0]
    if form_date:
        search_date = datetime.strptime(form_date,'%Y-%m-%d').date()

    pageination = LoadStatistic.query.filter_by(host_name = host_name, date_time = search_date).paginate(
        page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
        error_out = False
    )
    statistic_data = pageination.items
    if not statistic_data:
        last = LoadStatistic.query.filter_by(host_name = host_name).order_by(LoadStatistic.id.desc()).first()
        if last:
            pageination = LoadStatistic.query.filter_by(host_name = host_name, date_time = last.date_time).paginate(
                page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
                error_out = False
            )
            statistic_data = pageination.items

    statistic_data = [(index, item) for index, item in enumerate(statistic_data) ]
    return render_template('statistic.html',
        statistic_data = statistic_data,
        host_list = host_list,
        action = 'load_statistic',
        host_name = host_name,
        date = form_date,
        pageination = pageination)




@main.route('/load_statistic_host')
@login_required
def load_statistic_host():
    '''按设备统计用户数量'''

    city = session.get('city')
    load_statistic_host_data = []

    host_list = get_host(city)
    for i in host_list:
        data = LoadStatisticHost.query.filter_by(host_name = i).order_by(LoadStatisticHost.id.desc()).first()
        if data:
            load_statistic_host_data.append(data)

    load_statistic_host_data = [(index, item) for index, item in enumerate(load_statistic_host_data) ]
    return render_template('statistic_host.html',
        load_statistic_host_data = load_statistic_host_data, 
        action = 'load_statistic_host')

@main.route('/service_statistic')
@login_required
def service_statistic():
    '''端口业务类型统计'''

    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))
    host_list = get_host(session.get('city'))


    host_list = [(index, item) for index, item in enumerate(host_list) ]
    form = PortServiceForm()
    return render_template('report/service_statistic.html', host_list = host_list, action='service_statistic', form = form)



@main.route('/install_base')
@login_required
def install_base():
    '''installbase统计表'''

    install_base_data = InstallBase.query.filter_by(date_time = date.today()).all()
    if not install_base_data:
        last = InstallBase.query.order_by(InstallBase.id.desc()).first()
        if last:
            install_base_data = InstallBase.query.filter_by(date_time = last.date_time).all()

    install_base_data = [(index, item) for index, item in enumerate(install_base_data) ]
    return render_template('report/install_base.html',
        data = install_base_data,
        action = 'install_base')

@main.route('/net_flow')
@login_required
def net_flow():
    '''重要网络流量统计'''

    net_flow_data = []
    search_date = date.today()
    city_list = get_city_list()
    for i in city_list:
        host_list = get_host(i)
        for j in host_list:
            temp = NetFlow.query.filter_by(site_name = j, date_time = search_date).all()
            if not temp:
                last = NetFlow.query.filter_by(site_name=j).order_by(NetFlow.id.desc()).first()
                if last:
                    net_flow_data += NetFlow.query.filter_by(site_name = j, date_time = last.date_time).all()
            else:
                net_flow_data += temp

    return render_template('report/net_flow.html',
        data = net_flow_data,
        action = 'net_flow')

@main.route('/case_lib', methods= ['GET', 'POST'])
@login_required
def case_lib():
    '''典型案例'''

    page = request.args.get('page', 1, type=int)
    _from = request.args.get('from')
    form = CaseUploadForm()
    count = CaseLib.query.count()
    if count == 0:
        abort(404)

    if request.method == 'POST':
        search_case_lib_keyword = request.form.get('keyword')
        search_case_lib_date = request.form.get('date')

        if search_case_lib_keyword:
            session['search_case_lib_keyword'] = search_case_lib_keyword
        if search_case_lib_date:
            session['search_case_lib_date'] = datetime.strptime(search_case_lib_date,'%Y-%m-%d').date()

    case_lib = CaseLib.query
    if session.get('search_case_lib_keyword'):
        case_lib = case_lib.filter(CaseLib.describe.like("%" + session.get('search_case_lib_keyword') + "%"))
        session['search_case_lib_keyword'] = ''
    if session.get('search_case_lib_date'):
        case_lib = case_lib.filter(CaseLib.date_time == session.get('search_case_lib_date'))
        session['search_case_lib_date'] = ''


    pageination = case_lib.paginate(
        page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
        error_out = False
    )

    case_lib_data = pageination.items
    case_lib_data = [(index, item) for index, item in enumerate(case_lib_data) ]

    if _from == 'case_lib':
        url = 'case_lib/case_lib.html'
    else:
        url = 'case_lib/case_upload.html'
    return render_template(url,
        case_lib_data = case_lib_data,
        pageination = pageination,
        form=form)

@main.route('/case_delete/<id>')
@login_required
def case_delete(id):
    '''案例删除'''

    case = CaseLib.query.filter_by(id = id).first()
    save_path = os.path.join('app', 'static', 'caselib', case.file_name)
    os.remove(save_path)
    db.session.delete(case)
    db.session.commit()
    
    return redirect(url_for('main.case_upload'))

@main.route('/case_upload/nokia2020' , methods=['GET', 'POST'])
@login_required
def case_upload():
    '''案例上传'''

    page = request.args.get('page', 1, type=int)
    form = CaseUploadForm()
    if form.validate_on_submit():
        f = form.upload_file.data
        describe = form.describe.data
        # random_name = str(time.time()) + '.' +f.filename.split('.')[-1]
        random_name = f.filename
        f.save(os.path.join('app', 'static', 'caselib', random_name))
        # flash('上传成功')
        url = url_for('static', filename = os.path.join('caselib', random_name))
        url2 = unquote(url, encoding='utf-8')
        today = date.today()
        #上传信息入库
        case_lib = CaseLib(
            file_name = random_name,
            describe = describe,
            file_url = url2,
            date_time = today
        )

        db.session.add(case_lib)
        db.session.commit()

        return redirect(url_for('main.case_upload'))


    pageination = CaseLib.query.paginate(
        page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
        error_out = False
    )

    case_lib_data = pageination.items
    case_lib_data = [(index, item) for index, item in enumerate(case_lib_data) ]

    return render_template('case_lib/case_upload.html',
        form = form,
        case_lib_data = case_lib_data,
        pageination = pageination)

@main.route('/zuxun/<host_name>')
@login_required
def zuxun(host_name):
    '''组巡工具'''

    city = session.get('city')
    if not city:
        return redirect(url_for('main.city_list'))

    zuxun_data = ZuXun.query.filter_by(host_name = host_name, date_time = date.today()).all()
    if not zuxun_data:
        last = ZuXun.query.filter_by(host_name = host_name).order_by(ZuXun.id.desc()).first()
        if last:
            zuxun_data = ZuXun.query.filter_by(host_name = host_name, date_time = last.date_time).all()


    return render_template('zuxun/zuxun.html',
        zuxun_data = zuxun_data, 
        action = 'zuxun', 
        host_name=host_name)

@main.route('/zuxun_all/<host_name>', methods=['GET', 'POST'])
@login_required
def zuxun_all(host_name):
    '''组巡搜索'''

    host_list = get_host(session.get('city'))
    count = ZuXun.query.filter_by(date_time = date.today()).count()
    if count == 0:
        abort(404)

    if request.method == 'POST':
        session['search_zuxun_date'] = ''
        session['search_zuxun_host_name'] = ''

        date_str = request.form.get('date')
        host_name_str = request.form.get('host_name')

        if date_str:
            session['search_zuxun_date'] = datetime.strptime(date_str,'%Y-%m-%d').date()
        if host_name_str:
            session['search_zuxun_host_name'] = host_name_str
        else:
            session['search_zuxun_host_name'] = host_list[0]

    #搜索
    zuxun = ZuXun.query
    if session.get('search_zuxun_host_name'):
        zuxun = zuxun.filter(ZuXun.host_name == session.get('search_zuxun_host_name'))
    if session.get('search_zuxun_date'):
        zuxun = zuxun.filter(ZuXun.date_time == session.get('search_zuxun_date'))
    else:
        zuxun = zuxun.filter(ZuXun.date_time == date.today())

    zuxun_data = zuxun.all()

    return render_template('zuxun/zuxun_all.html',
        zuxun_data = zuxun_data,
        host_list = host_list,
        action = 'zuxun_all',
        host_name = host_name)

@main.route('/model', methods=['GET', 'POST'])
@login_required
def model():
    '''脚本自动生成-模板列表'''

    page = request.args.get('page', 1, type=int)

    generate_config = GenerateConfig.query
    generate_config = generate_config.filter(GenerateConfig.model_type.in_([0,1]))

    pageination = generate_config.paginate(
        page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
        error_out = False
    )

    model_data = pageination.items
    model_data = [(index, item) for index, item in enumerate(model_data) ]
    return render_template('generate_config/model.html',
        model_data = model_data,
        pageination = pageination)

@main.route('/model_list', methods=['GET', 'POST'])
@login_required
def model_list():
    '''脚本自动生成-模板组列表'''

    page = request.args.get('page', 1, type=int)

    generate_config = GenerateConfig.query
    generate_config = generate_config.filter(GenerateConfig.model_type == 2)

    pageination = generate_config.paginate(
        page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
        error_out = False
    )

    model_data = pageination.items
    model_data = [(index, item) for index, item in enumerate(model_data) ]
    return render_template('generate_config/model.html',
        model_data = model_data,
        pageination = pageination)

@main.route('/model_view/<id>')
@login_required
def model_view(id):
    '''模板查看'''
    model_data = GenerateConfig.query.filter_by(id=id).first()
    return render_template('generate_config/model_view.html',
        model_data = model_data)

@main.route('/model_delete/<id>')
@login_required
def model_delete(id):
    '''模板删除'''

    model_data = GenerateConfig.query.filter_by(id=id).first()
    #内置模板无法删除
    if model_data.model_type == 0:
        flash('内置模板无法删除')
        return redirect(url_for('main.model'))
    db.session.delete(model_data)
    db.session.commit()

    if model_data.model_type == 2:
        return redirect(url_for('main.model_list'))
    else:
        return redirect(url_for('main.model'))

    # return redirect(url_for('main.model'))

@main.route('/model_list_delete/<id>')
@login_required
def model_list_delete(id):
    '''模板组删除'''

    model_data = GenerateConfig.query.filter_by(id=id).first()
    db.session.delete(model_data)
    db.session.commit()

    return redirect(url_for('main.model_list'))

@main.route('/model_modify', methods=['GET', 'POST'])
@login_required
def model_modify():
    '''模板修改'''

    model_data = None
    form = ModelForm()

    if form.validate_on_submit():
        #模板名称不能重复
        model_from_name = GenerateConfig.query.filter_by(name=form.name.data).first()
        if model_from_name and model_from_name.id != int(form.id.data):
            flash('模板名称已存在，请重新输入')
            return redirect(url_for('main.model_modify', id=form.id.data))

        model_data = GenerateConfig.query.filter_by(id=form.id.data).first()
        model_data.name = form.name.data
        model_data.content = form.content.data

        db.session.add(model_data)
        db.session.commit()
        if model_data.model_type == 2:
            return redirect(url_for('main.model_list'))
        else:
            return redirect(url_for('main.model'))


    id = request.args.get('id', type=int)
    if id:
        model_data = GenerateConfig.query.filter_by(id=id).first()
        if model_data.model_type == 0:
            flash('内置模板无法修改')
            return redirect(url_for('main.model'))
        form.id.data = id
        form.name.data = model_data.name
        form.content.data = model_data.content

    return render_template('generate_config/model_modify.html',
        form = form)


@main.route('/model_save_as', methods=['GET', 'POST'])
@login_required
def model_save_as():
    '''模板另存'''

    model_data = None
    form = ModelForm()

    if form.validate_on_submit():
        #模板名称不能重复
        model_from_name = GenerateConfig.query.filter_by(name=form.name.data).first()
        if model_from_name:
            flash('模板名称已存在，请重新输入')
            return redirect(url_for('main.model_save_as', id=form.id.data))
        model_data = GenerateConfig(
            name = form.name.data,
            content = form.content.data,
            model_type = 1,
            date_time = date.today()
        )
        db.session.add(model_data)
        db.session.commit()
        return redirect(url_for('main.model'))

    id = request.args.get('id', type=int)
    if id:
        model_data = GenerateConfig.query.filter_by(id=id).first()
        form.id.data = id
        form.name.data = model_data.name
        form.content.data = model_data.content

    return render_template('generate_config/model_save_as.html',
        form = form)

@main.route('/create_model_list', methods=['GET', 'POST'])
@login_required
def create_model_list():
    '''新建模板组'''

    form = ModelListCreateForm()
    model_data = GenerateConfig.query.filter_by(model_type=0).all()
    form.model_names.choices = [(item.id, item.name) for item in model_data]

    if form.validate_on_submit():
         #模板名称不能重复
        model_from_name = GenerateConfig.query.filter_by(name=form.name.data).first()
        if model_from_name:
            flash('模板名称已存在，请重新输入')
            return redirect(url_for('main.create_model_list'))
        model_str = ''
        for i in form.model_names.data:
            model_data = GenerateConfig.query.filter_by(id=i).first()
            model_str += model_data.content + '\n'

        model_data = GenerateConfig(
            name = form.name.data,
            content = model_str,
            model_type = 2,
            date_time = date.today()
        )
        db.session.add(model_data)
        db.session.commit()
        return redirect(url_for('main.model_list'))

    return render_template('generate_config/model_list_create.html',
        form = form)

@main.route('/model_select', methods=['GET', 'POST'])
@login_required
def model_select():
    '''模板选择'''

    form = ModelSelectForm()
    generate_config = GenerateConfig.query
    generate_config = generate_config.filter(GenerateConfig.model_type.in_([1,2]))
    custom_model_data = generate_config.all()
    form.custom_model_names.choices = [(item.id, item.name) for item in custom_model_data]
    model_data = GenerateConfig.query.filter_by(model_type = 0).all()
    form.model_names.choices = [(item.id, item.name) for item in model_data]

    if form.validate_on_submit():
        model_str = ''
        for i in form.model_names.data:
            model_data = GenerateConfig.query.filter_by(id=i).first()
            model_str += model_data.content + '\n'

        for i in form.custom_model_names.data:
            model_data = GenerateConfig.query.filter_by(id=i).first()
            model_str += model_data.content + '\n'

        if model_str == '':
            flash('请选择模板或模板组')
            return redirect(url_for('main.model_select', form=form))
        else:
            session['model_str'] = model_str
        titles = parse_model(model_str)
        return render_template('generate_config/enter_value.html',
        titles = titles)

    return render_template('generate_config/select_model_form.html',
        form = form)

@main.route('/make_config', methods=['POST'])
@login_required
def make_config():
    '''生成配置'''
    
    inputs = list(request.form.items(True))
    model_str = session.get('model_str')
    model_data = generate_config(model_str, inputs)
    session['model_data'] = model_data
    return render_template('generate_config/result.html',
        model_data = model_data)

@main.route('/download_model_data')
@login_required
def download_model_data():
    '''保存生成的配置'''

    model_data = session.get('model_data')
    with open(os.path.join('app','static', 'config.log'), 'w') as f:
        f.write(model_data)

    url = url_for('static', filename = 'config.log')
    url2 = unquote(url, encoding='utf-8')
    return redirect(url2)

@main.route('/generate_config_instruction')
@login_required
def generate_config_instruction():
    '''脚本自动生成工具使用说明'''

    return render_template('generate_config/instruction.html')


@main.route('/report_port_search/<report_name>', methods=['GET', 'POST'])
@login_required
def report_port_search(report_name):
    '''报表搜索'''

    page = request.args.get('page', 1, type=int)
    if request.method == 'POST':
        session['search_port_detail_port'] = ''
        session['search_port_detail_date'] = ''
        session['search_port_detail_host_name'] = ''

        session['search_port_statistic_host_name'] = ''
        session['search_port_statistic_host_ip'] = ''
        session['search_port_statistic_date'] = ''

        session['search_card_detail_host_name'] = ''
        session['search_card_detail_host_ip'] = ''
        session['search_card_detail_date'] = ''

        session['search_card_statistic_host_name'] = ''
        session['search_card_statistic_host_ip'] = ''
        session['search_card_statistic_date'] = ''

        session['search_mda_detail_host_name'] = ''
        session['search_mda_detail_host_ip'] = ''
        session['search_mda_detail_date'] = ''

        session['search_mda_statistic_host_name'] = ''
        session['search_mda_statistic_host_ip'] = ''
        session['search_mda_statistic_date'] = ''

        session['search_address_collect_host_name'] = ''
        session['search_address_collect_host_ip'] = ''
        session['search_address_collect_date'] = ''


        search_port_detail_port = request.form.get('port')
        search_port_detail_date = request.form.get('date')
        search_port_detail_host_name = request.form.get('host_name')

        search_port_statistic_host_ip = request.form.get('port_statistic_host_ip')
        search_port_statistic_date = request.form.get('port_statistic_date')
        search_port_statistic_host_name = request.form.get('port_statistic_host_name')

        search_card_detail_host_ip = request.form.get('card_detail_host_ip')
        search_card_detail_date = request.form.get('card_detail_date')
        search_card_detail_host_name = request.form.get('card_detail_host_name')

        search_card_statistic_host_ip = request.form.get('card_statistic_host_ip')
        search_card_statistic_date = request.form.get('card_statistic_date')
        search_card_statistic_host_name = request.form.get('card_statistic_host_name')

        search_mda_detail_host_ip = request.form.get('mda_detail_host_ip')
        search_mda_detail_date = request.form.get('mda_detail_date')
        search_mda_detail_host_name = request.form.get('mda_detail_host_name')

        search_mda_statistic_host_ip = request.form.get('mda_statistic_host_ip')
        search_mda_statistic_date = request.form.get('mda_statistic_date')
        search_mda_statistic_host_name = request.form.get('mda_statistic_host_name')

        search_address_collect_host_ip = request.form.get('address_collect_host_ip')
        search_address_collect_date = request.form.get('address_collect_date')
        search_address_collect_host_name = request.form.get('address_collect_host_name')

        if search_port_detail_port:
            session['search_port_detail_port'] = search_port_detail_port
        if search_port_detail_host_name:
            session['search_port_detail_host_name'] = search_port_detail_host_name
        if search_port_detail_date:
            session['search_port_detail_date'] = datetime.strptime(search_port_detail_date,'%Y-%m-%d').date()

        if search_port_statistic_host_ip:
            session['search_port_statistic_host_ip'] = search_port_statistic_host_ip
        if search_port_statistic_date:
            session['search_port_statistic_date'] = datetime.strptime(search_port_statistic_date,'%Y-%m-%d').date()
        if search_port_statistic_host_name:
            session['search_port_statistic_host_name'] = search_port_statistic_host_name

        if search_card_detail_host_ip:
            session['search_card_detail_host_ip'] = search_card_detail_host_ip
        if search_card_detail_date:
            session['search_card_detail_date'] = datetime.strptime(search_card_detail_date,'%Y-%m-%d').date()
        if search_card_detail_host_name:
            session['search_card_detail_host_name'] = search_card_detail_host_name

        if search_card_statistic_host_ip:
            session['search_card_statistic_host_ip'] = search_card_statistic_host_ip
        if search_card_statistic_date:
            session['search_card_statistic_date'] = datetime.strptime(search_card_statistic_date,'%Y-%m-%d').date()
        if search_card_statistic_host_name:
            session['search_card_statistic_host_name'] = search_card_statistic_host_name

        if search_mda_detail_host_ip:
            session['search_mda_detail_host_ip'] = search_mda_detail_host_ip
        if search_mda_detail_date:
            session['search_mda_detail_date'] = datetime.strptime(search_mda_detail_date,'%Y-%m-%d').date()
        if search_mda_detail_host_name:
            session['search_mda_detail_host_name'] = search_mda_detail_host_name

        if search_mda_statistic_host_ip:
            session['search_mda_statistic_host_ip'] = search_mda_statistic_host_ip
        if search_mda_statistic_date:
            session['search_mda_statistic_date'] = datetime.strptime(search_mda_statistic_date,'%Y-%m-%d').date()
        if search_mda_statistic_host_name:
            session['search_mda_statistic_host_name'] = search_mda_statistic_host_name

        if search_address_collect_host_ip:
            session['search_address_collect_host_ip'] = search_address_collect_host_ip
        if search_address_collect_date:
            session['search_address_collect_date'] = datetime.strptime(search_address_collect_date,'%Y-%m-%d').date()
        if search_address_collect_host_name:
            session['search_address_collect_host_name'] = search_address_collect_host_name

    #端口明细
    if report_name == 'port_detail':
        count = CardPort1.query.filter_by(date_time = date.today()).count()
        if count == 0:
            abort(404)

        #搜索
        cardport1 = CardPort1.query
        if session.get('search_port_detail_port'):
            cardport1 = cardport1.filter(CardPort1.port == session.get('search_port_detail_port'))
        if session.get('search_port_detail_host_name'):
            cardport1 = cardport1.filter(CardPort1.host_name == session.get('search_port_detail_host_name'))
        if session.get('search_port_detail_date'):
            a = session.get('search_port_detail_date')
            cardport1 = cardport1.filter(CardPort1.date_time == session.get('search_port_detail_date'))
        else:
            cardport1 = cardport1.filter(CardPort1.date_time == date.today())

        pageination = cardport1.paginate(
            page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
            error_out = False
        )
        card_data = pageination.items

        return render_template('search/report.html', 
            card_data = card_data, 
            report_name = 'port_detail', 
            pageination = pageination)

    #端口统计
    elif report_name == 'port_statistic':
        count = PortStatistic.query.filter_by(date_time = date.today()).count()
        if count == 0:
            abort(404)

         #搜索
        port_statistic = PortStatistic.query
        if session.get('search_port_statistic_host_ip'):
            port_statistic = port_statistic.filter(PortStatistic.host_ip == session.get('search_port_statistic_host_ip'))
        if session.get('search_port_statistic_host_name'):
            port_statistic = port_statistic.filter(PortStatistic.host_name == session.get('search_port_statistic_host_name'))
        if session.get('search_port_statistic_date'):
            port_statistic = port_statistic.filter(PortStatistic.date_time == session.get('search_port_statistic_date'))
        else:
            port_statistic = port_statistic.filter(PortStatistic.date_time == date.today())


        pageination = port_statistic.paginate(
            page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
            error_out = False
        )
        port_statistic_data = pageination.items

        return render_template('search/port_statistic.html', 
            port_statistic_data = port_statistic_data,
            report_name = 'port_statistic', 
            pageination = pageination)

    #card明细
    elif report_name == 'card_detail':
        count = CardDetail.query.filter_by(date_time = date.today()).count()
        if count == 0:
            abort(404)

         #搜索
        card_detail = CardDetail.query
        if session.get('search_card_detail_host_ip'):
            card_detail = card_detail.filter(CardDetail.host_ip == session.get('search_card_detail_host_ip'))
        if session.get('search_card_detail_host_name'):
            card_detail = card_detail.filter(CardDetail.host_name == session.get('search_card_detail_host_name'))
        if session.get('search_card_detail_date'):
            card_detail = card_detail.filter(CardDetail.date_time == session.get('search_card_detail_date'))
        else:
            card_detail = card_detail.filter(CardDetail.date_time == date.today())


        pageination = card_detail.paginate(
            page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
            error_out = False
        )
        card_detail_data = pageination.items

        return render_template('search/card_detail.html', 
            card_detail_data = card_detail_data,
            report_name = 'card_detail', 
            pageination = pageination)

    #card统计
    elif report_name == 'card_statistic':
        count = CardStatistic.query.filter_by(date_time = date.today()).count()
        if count == 0:
            abort(404)

        #搜索
        card_statistic = CardStatistic.query
        if session.get('search_card_statistic_host_ip'):
            card_statistic = card_statistic.filter(CardStatistic.host_ip == session.get('search_card_statistic_host_ip'))
        if session.get('search_card_statistic_host_name'):
            card_statistic = card_statistic.filter(CardStatistic.host_name == session.get('search_card_statistic_host_name'))
        if session.get('search_card_statistic_date'):
            card_statistic = card_statistic.filter(CardStatistic.date_time == session.get('search_card_statistic_date'))
        else:
            card_statistic = card_statistic.filter(CardStatistic.date_time == date.today())


        pageination = card_statistic.paginate(
            page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
            error_out = False
        )
        card_statistic_data = pageination.items

        return render_template('search/card_statistic.html', 
            card_statistic_data = card_statistic_data,
            report_name = 'card_statistic', 
            pageination = pageination)

    #mda明细
    elif report_name == 'mda_detail':

        count = MdaDetail.query.filter_by(date_time = date.today()).count()
        if count == 0:
            abort(404)

         #搜索
        mda_detail = MdaDetail.query
        if session.get('search_mda_detail_host_ip'):
            mda_detail = mda_detail.filter(MdaDetail.host_ip == session.get('search_mda_detail_host_ip'))
        if session.get('search_mda_detail_host_name'):
            mda_detail = mda_detail.filter(MdaDetail.host_name == session.get('search_mda_detail_host_name'))
        if session.get('search_mda_detail_date'):
            mda_detail = mda_detail.filter(MdaDetail.date_time == session.get('search_mda_detail_date'))
        else:
            mda_detail = mda_detail.filter(MdaDetail.date_time == date.today())


        pageination = mda_detail.paginate(
            page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
            error_out = False
        )
        mda_detail_data = pageination.items

        return render_template('search/mda_detail.html', 
            mda_detail_data = mda_detail_data,
            report_name = 'mda_detail', 
            pageination = pageination)

    #mda统计
    elif report_name == 'mda_statistic':
        
        count = MdaStatistic.query.filter_by(date_time = date.today()).count()
        if count == 0:
            abort(404)

        #搜索
        mda_statistic = MdaStatistic.query
        if session.get('search_mda_statistic_host_ip'):
            mda_statistic = mda_statistic.filter(MdaStatistic.host_ip == session.get('search_mda_statistic_host_ip'))
        if session.get('search_mda_statistic_host_name'):
            mda_statistic = mda_statistic.filter(MdaStatistic.host_name == session.get('search_mda_statistic_host_name'))
        if session.get('search_mda_statistic_date'):
            mda_statistic = mda_statistic.filter(MdaStatistic.date_time == session.get('search_mda_statistic_date'))
        else:
            mda_statistic = mda_statistic.filter(MdaStatistic.date_time == date.today())


        pageination = mda_statistic.paginate(
            page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
            error_out = False
        )
        mda_statistic_data = pageination.items

        return render_template('search/mda_statistic.html', 
            mda_statistic_data = mda_statistic_data,
            report_name = 'mda_statistic', 
            pageination = pageination)

    #地址采集
    elif report_name == 'address_collect':
        
        count = AddressCollect.query.filter_by(date_time = date.today()).count()
        if count == 0:
            abort(404)

        #搜索
        address_collect = AddressCollect.query
        if session.get('search_address_collect_host_ip'):
            address_collect = address_collect.filter(AddressCollect.host_ip == session.get('search_address_collect_host_ip'))
        if session.get('search_address_collect_host_name'):
            address_collect = address_collect.filter(AddressCollect.host_name == session.get('search_address_collect_host_name'))
        if session.get('search_address_collect_date'):
            address_collect = address_collect.filter(AddressCollect.date_time == session.get('search_address_collect_date'))
        else:
            address_collect = address_collect.filter(AddressCollect.date_time == date.today())


        pageination = address_collect.paginate(
            page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
            error_out = False
        )
        address_collect_data = pageination.items

        return render_template('search/address_collect.html', 
            address_data = address_collect_data,
            report_name = 'address_collect', 
            pageination = pageination)

@main.route('/load_statistic_search/<load_name>', methods=['GET', 'POST'])
@login_required
def load_statistic_search(load_name):
    '''业务负载搜索'''

    page = request.args.get('page', 1, type=int)
    if request.method == 'POST':
        session['search_load_statistic_host_ip'] = ''
        session['search_load_statistic_date'] = ''
        session['search_load_statistic_host_name'] = ''

        session['search_load_statistic_host_host_name'] = ''
        session['search_load_statistic_host_host_ip'] = ''
        session['search_load_statistic_host_date'] = ''

        search_load_statistic_host_ip = request.form.get('load_statistic_host_ip')
        search_load_statistic_date = request.form.get('load_statistic_date')
        search_load_statistic_host_name = request.form.get('load_statistic_host_name')

        search_load_statistic_host_host_ip = request.form.get('load_statistic_host_host_ip')
        search_load_statistic_host_date = request.form.get('load_statistic_host_date')
        search_load_statistic_host_host_name = request.form.get('load_statistic_host_host_name')

        if search_load_statistic_host_ip:
            session['search_load_statistic_host_ip'] = search_load_statistic_host_ip
        if search_load_statistic_host_name:
            session['search_load_statistic_host_name'] = search_load_statistic_host_name
        if search_load_statistic_date:
            session['search_load_statistic_date'] = datetime.strptime(search_load_statistic_date,'%Y-%m-%d').date()

        if search_load_statistic_host_host_ip:
            session['search_load_statistic_host_host_ip'] = search_load_statistic_host_host_ip
        if search_load_statistic_host_date:
            session['search_load_statistic_host_date'] = datetime.strptime(search_load_statistic_host_date,'%Y-%m-%d').date()
        if search_load_statistic_host_host_name:
            session['search_load_statistic_host_host_name'] = search_load_statistic_host_host_name


    if load_name == 'load_statistic':
    
        count = LoadStatistic.query.filter_by(date_time = date.today()).count()
        if count == 0:
            abort(404)

        #搜索
        load_statistic = LoadStatistic.query
        if session.get('search_load_statistic_host_ip'):
            load_statistic = load_statistic.filter(LoadStatistic.host_ip == session.get('search_load_statistic_host_ip'))
        if session.get('search_load_statistic_host_name'):
            load_statistic = load_statistic.filter(LoadStatistic.host_name == session.get('search_load_statistic_host_name'))
        if session.get('search_load_statistic_date'):
            load_statistic = load_statistic.filter(LoadStatistic.date_time == session.get('search_load_statistic_date'))
        else:
            load_statistic = load_statistic.filter(LoadStatistic.date_time == date.today())


        pageination = load_statistic.paginate(
            page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
            error_out = False
        )
        load_statistic_data = pageination.items

        return render_template('search/statistic.html', 
            load_statistic_data = load_statistic_data,
            load_name = 'load_statistic', 
            pageination = pageination)
    else:
        #业务负载统计，按设备统计
        count = LoadStatisticHost.query.filter_by(date_time = date.today()).count()
        if count == 0:
            abort(404)

        #搜索
        load_statistic_host = LoadStatisticHost.query
        if session.get('search_load_statistic_host_host_ip'):
            load_statistic_host = load_statistic_host.filter(LoadStatisticHost.host_ip == session.get('search_load_statistic_host_host_ip'))
        if session.get('search_load_statistic_host_host_name'):
            a = session.get('search_load_statistic_host_host_name')
            load_statistic_host = load_statistic_host.filter(LoadStatisticHost.host_name == session.get('search_load_statistic_host_host_name'))
        if session.get('search_load_statistic_host_date'):
            load_statistic_host = load_statistic_host.filter(LoadStatisticHost.date_time == session.get('search_load_statistic_host_date'))
        else:
            load_statistic_host = load_statistic_host.filter(LoadStatisticHost.date_time == date.today())


        pageination = load_statistic_host.paginate(
            page, per_page = current_app.config['FLASKY_POSTS_PER_PAGE'],
            error_out = False
        )
        load_statistic_host_data = pageination.items

        return render_template('search/statistic_host.html', 
            load_statistic_host_data = load_statistic_host_data,
            load_name = 'load_statistic_host', 
            pageination = pageination)

@main.route('/xunjian_search/<xunjian_name>', methods= ['GET', 'POST'])
@login_required
def xunjian_search(xunjian_name):
    '''巡检搜索'''

    host_list = get_host(session.get('city'))
    page = request.args.get('page', 1, type=int)
    count = XunJian.query.filter_by(date_time = date.today()).count()
    if count == 0:
        abort(404)

    if request.method == 'POST':
        session['search_xunjian_date'] = ''
        session['search_xunjian_host_name'] = ''

        session['search_xunjian_all_date'] = ''
        session['search_xunjian_all_host_name'] = ''

        session['search_xunjian_log_date'] = ''
        session['search_xunjian_log_host_name'] = ''

        search_xunjian_date = request.form.get('xunjian_date')
        search_xunjian_host_name = request.form.get('xunjian_host_name')

        search_xunjian_all_date = request.form.get('xunjian_all_date')
        search_xunjian_all_host_name = request.form.get('xunjian_all_host_name')

        search_xunjian_log_date = request.form.get('xunjian_log_date')
        search_xunjian_log_host_name = request.form.get('xunjian_log_host_name')

        if search_xunjian_date:
            session['search_xunjian_date'] = search_xunjian_date
        if search_xunjian_host_name:
            session['search_xunjian_host_name'] = search_xunjian_host_name
        else:
            session['search_xunjian_host_name'] = host_list[0]

        if search_xunjian_all_date:
            session['search_xunjian_all_date'] = search_xunjian_all_date
        if search_xunjian_all_host_name:
            session['search_xunjian_all_host_name'] = search_xunjian_all_host_name

        if search_xunjian_log_date:
            session['search_xunjian_log_date'] = search_xunjian_log_date
        if search_xunjian_log_host_name:
            session['search_xunjian_log_host_name'] = search_xunjian_log_host_name

    if xunjian_name == 'xunjian':
        #设备巡检报告

        #搜索
        xunjian = XunJian.query
        if session.get('search_xunjian_host_name'):
            xunjian = xunjian.filter(XunJian.host_name == session.get('search_xunjian_host_name'))
        else:
            xunjian = xunjian.filter(XunJian.host_name == host_list[0])

        if session.get('search_xunjian_date'):
            xunjian = xunjian.filter(XunJian.date_time == session.get('search_xunjian_date'))
        else:
            xunjian = xunjian.filter(XunJian.date_time == date.today())

        xunjian_data = xunjian.all()

        return render_template('search/xunjian.html', 
            xunjian_data = xunjian_data,
            xunjian_name = 'xunjian',
            host_list = host_list)

    elif xunjian_name == 'xunjian_all':
        #巡检输出

        #搜索
        xunjian = XunJian.query
        if session.get('search_xunjian_all_host_name'):
            xunjian = xunjian.filter(XunJian.host_name == session.get('search_xunjian_all_host_name'))
        else:
            xunjian = xunjian.filter(XunJian.host_name == host_list[0])
        if session.get('search_xunjian_all_date'):
            xunjian = xunjian.filter(XunJian.date_time == session.get('search_xunjian_all_date'))
        else:
            xunjian = xunjian.filter(XunJian.date_time == date.today())

        xunjian_data = xunjian.all()

        return render_template('search/xunjian_all.html',
            xunjian_all_data = xunjian_data,
            xunjian_name = 'xunjian_all',
            host_list = host_list)

    else:
        #全量输出
        city = ''
        host_name = ''
        date_str = ''
        log_str = ''

        city = session.get('city')
        host_name = session.get('search_xunjian_log_host_name')
        date_str = session.get('search_xunjian_log_date')

        if city and host_name and date_str:
            log_str = get_log_from_date(city, host_name, date_str)

        return render_template('search/xunjian_log.html', 
            log_str = log_str,
            xunjian_name = 'xunjian_log',
            host_list = host_list)
        
    
    return '巡检搜索'

@main.route('/save_db')
def save_db():
    '''把所有数据入库'''

    save_city = request.args.get('city')
    citys = get_city_list()
    for i in citys:
        if save_city and i != save_city:
            continue
        hosts = get_host(i)
        # city_count = CardPort1.query.filter_by(city = i, date_time = date.today()).count()
        # if city_count > 0:
        #     logging.info('city {} today is saved'.format(i))
        #     continue
        for j in hosts:
            today_log = get_log(i, j)
            yesterday_log = get_log(i, j, 1)
            if today_log:
                #端口明细
                save_port_detail(i, j, today_log)
                # t_port_detail = threading.Thread(target=save_port_detail, args=(i, j, today_log))
                # t_port_detail.start()
                #端口统计
                save_port_statistic(i, j, today_log)
                # t_port_statistic = threading.Thread(target=save_port_statistic, args=(i, j, today_log))
                # t_port_statistic.start()
                #card明细
                save_card_detail(i, j, today_log)
                # t_card_detail = threading.Thread(target=save_card_detail, args=(i, j, today_log))
                # t_card_detail.start()
                #card统计
                save_card_statistic(i, j, today_log)
                # t_card_statistic = threading.Thread(target=save_card_statistic, args=(i, j, today_log))
                # t_card_statistic.start()
                #mda明细
                save_mda_detail(i, j, today_log)
                # t_mda_detail = threading.Thread(target=save_mda_detail, args=(i, j, today_log))
                # t_mda_detail.start()
                # mda统计
                save_mda_statistic(i, j, today_log)
                # t_mda_statistic = threading.Thread(target=save_mda_statistic, args=(i, j, today_log))
                # t_mda_statistic.start()
                #业务负载统计-按端口统计
                save_load_statistic(i, j, today_log)
                # t_load_statistic = threading.Thread(target=save_load_statistic, args=(i, j, today_log))
                # t_load_statistic.start()
                #业务负载统计-按设备统计
                save_load_statistic_host(i, j, today_log)
                # t_load_statistic_host = threading.Thread(target=save_load_statistic_host, args=(i, j, today_log))
                # t_load_statistic_host.start()
                #端口业务类型统计
                # save_service_statistic(i, j, today_log)
                #地址采集
                save_address_collect(i, j, today_log)
                # t_address_collect = threading.Thread(target=save_address_collect, args=(i, j, today_log))
                # t_address_collect.start()
                #组巡
                save_zuxun(i, j, today_log)
                # t_zuxun = threading.Thread(target=save_zuxun, args=(i, j, today_log))
                # t_zuxun.start()
                #重要网络流量
                save_netflow(i, j, today_log)
                # t_netflow = threading.Thread(target=save_netflow, args=(i, j, today_log))
                # t_netflow.start()
                #专线统计
                save_special_line(i, j, today_log)
                #Nat Port 利用率统计
                save_nat_port(i, j, today_log)
            if today_log and yesterday_log:
                save_xunjian(i, j, today_log, yesterday_log)
                # t_xunjian = threading.Thread(target=save_xunjian, args=(i, j, today_log, yesterday_log))
                # t_xunjian.start()
                #巡检报告
                xj_report_all_host(i)
            else:
                if not today_log:
                    logging.info('host: {} not found today log'.format(j))
                if not yesterday_log:
                    logging.info('host: {} not found yesterday log'.format(j))
        #设备清单
        save_host_list(i)
        # t_host_list = threading.Thread(target=save_host_list, args=(i,))
        # t_host_list.start()
    #installbase统计
    save_install_base()
    #生成巡检报告
    xj_report_summary()
    # t_install_base = threading.Thread(target=save_install_base)
    # t_install_base.start()

    # db.session.remove()
    return '数据保存成功'


def save_xunjian(city, host_name, config_new, config_old):
    '''巡检数据入库'''

    # app = create_app('production')
    # app.app_context().push()
    today_data_count = XunJian.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('xunjian host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    xunjian_data = mobile.xunjian(config_new, config_old)
    if xunjian_data:
        logging.info('host: {} {} begin save, date: {}'.format(host_name, 'xunjian_data', date.today().strftime('%Y-%m-%d')))

    today = date.today()
    for msg, err, check_item in xunjian_data:
        xunjian = XunJian(
            city = city,
            host_name = host_name,
            check_item = check_item,
            err = err,
            msg = msg,
            date_time = today
        )

        db.session.add(xunjian)
    db.session.commit()
    db.session.close()

def save_port_detail(city, host_name, config):
    '''端口明细入库'''

    # app = create_app('production')
    # app.app_context().push()
    today_data_count = CardPort1.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('port_detail host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    ip = get_ip(config)
    report_data = get_report_data(config)
    ggl_data = get_port_ggl(config)
    ggl_port = []
    res = []

    i = 0
    for item in report_data:
        if item[-1] != 'MDX GIGE-T ' and item[-1] != 'MDI GIGE-T ':
            ggl_port.append(item[0])
            port_is_ok = '否'

            #判断端口是否有异常
            if item[2] == 'Up' and item[3] == 'No' and item[4] != 'Up':
                port_is_ok = '是'

            #判断端口ge或者10ge
            dk = ''
            if item[-1].startswith('10'):
                dk = '10GE'
            elif item[-1].startswith('GIGE'):
                dk = 'GE'
            else:
                dk = ''
            try:
                if item[2] == 'Up' and item[3] == 'Yes' and item[4] == 'Up':
                    if ggl_data[i][7] and (float(ggl_data[i][5]) < float(ggl_data[i][8]) or float(ggl_data[i][5]) > float(ggl_data[i][7])):
                        port_is_ok = '是'
                    if ggl_data[i][3] and (float(ggl_data[i][0]) < float(ggl_data[i][3]) or float(ggl_data[i][0]) > float(ggl_data[i][2])):
                        port_is_ok = '是'

                res.append(item + [ggl_data[i][0], ggl_data[i][5], ggl_data[i][2]+'|'+ ggl_data[i][3], ggl_data[i][7]+'|'+ ggl_data[i][8], ip, port_is_ok, dk])

            except Exception:
                continue
            
            i += 1
        else:
            res.append(item + ['','','','', ip, '', ''])


    if res:
        logging.info('host: {} {} begin save'.format(host_name, 'port_detail'))

    today = date.today()
    for item in res:

        #存入数据库
        card_port = CardPort1(
            city = city,
            host_name = host_name,
            host_ip = item[16],
            port = item[1],
            port_dk = item[18],
            admin_state = item[2],
            link_state = item[3],
            port_state = item[4],
            cfg_mtu = item[5],
            oper_mtu = item[6],
            lag = item[7],
            port_mode = item[8],
            port_encp = item[9],
            port_type = item[10],
            c_qs_s_xfp_mdimdx = item[11],
            optical_power = item[12],
            output_power = item[13],
            optical_warn = item[14],
            output_warn = item[15],
            is_abnormal = item[17],
            date_time = today
        )

        db.session.add(card_port)

    try:
        db.session.commit()
    except:
        db.session.rollback()
    db.session.close()

def save_port_statistic(city, host_name, config):
    '''端口统计入库'''

    # app = create_app('production')
    # app.app_context().push()
    today_data_count = PortStatistic.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('port_statistic host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    port_statistic_data = get_port_statistic(config)
    if port_statistic_data:
        logging.info('host: {} {} begin save'.format(host_name, 'port_statistic_data'))

    today = date.today()
    for item in port_statistic_data:

        #存入数据库
        port_statistic = PortStatistic(
            city = city,
            host_name = host_name,
            host_ip = item[1],
            port_type = item[2],
            port_num = item[3],
            used_num = item[4],
            unused_num = item[5],
            date_time = today
        )

        db.session.add(port_statistic)
    try:
        db.session.commit()
    except:
        db.session.rollback()
    db.session.close()
def save_card_detail(city, host_name, config):
    '''card明细入库'''

    # app = create_app('production')
    # app.app_context().push()
    today_data_count = CardDetail.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('card_detail host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    card_detail_data = get_card_detail(config)
    if card_detail_data:
        logging.info('host: {} {} begin save'.format(host_name, 'card_detail_data'))

    today = date.today()
    for item in card_detail_data:
        #存入数据库
        card_detail = CardDetail(
            city = city,
            host_name = host_name,
            host_ip = item[1],
            slot = item[2],
            card_type = item[3],
            admin_state = item[4],
            operational_state = item[5],
            serial_number = item[6],
            time_of_last_boot = item[7],
            temperature = item[8],
            temperature_threshold = item[9],
            is_abnormal = item[10],
            date_time = today
        )

        db.session.add(card_detail)
    db.session.commit()
    db.session.close()

def save_card_statistic(city, host_name, config):
    '''card统计入库'''

    # app = create_app('production')
    # app.app_context().push()
    today_data_count = CardStatistic.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('card_statistic host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    card_statistic_data = get_card_statistic(config)
    if card_statistic_data:
        logging.info('host: {} {} begin save'.format(host_name, 'card_statistic_data'))

    today = date.today()
    for item in card_statistic_data:

        #存入数据库
        card_statistic = CardStatistic(
            city = city,
            host_name = host_name,
            host_ip = item[1],
            card_type = item[2],
            card_num = item[3],
            date_time = today
        )

        db.session.add(card_statistic)
    db.session.commit()
    db.session.close()

def save_mda_detail(city, host_name, config):
    '''mda明细入库'''

    # app = create_app('production')
    # app.app_context().push()
    today_data_count = MdaDetail.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('mda_detail host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    mda_detail_data = get_mda_detail(config)
    if mda_detail_data:
        logging.info('host: {} {} begin save'.format(host_name, 'mda_detail_data'))

    today = date.today()
    for item in mda_detail_data:
        #存入数据库
        mda_detail = MdaDetail(
            city = city,
            host_name = host_name,
            host_ip = item[1],
            slot = item[2],
            mda = item[3],
            equipped_type = item[4],
            admin_state = item[5],
            operational_state = item[6],
            serial_number = item[7],
            time_of_last_boot = item[8],
            temperature = item[9],
            temperature_threshold = item[10],
            is_abnormal = item[11],
            date_time = today
        )

        db.session.add(mda_detail)
    db.session.commit()
    db.session.close()

def save_mda_statistic(city, host_name, config):
    '''mda统计入库'''

    # app = create_app('production')
    # app.app_context().push()
    today_data_count = MdaStatistic.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('mda_statistic host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    mda_statistic_data = get_mda_statistic(config)
    if mda_statistic_data:
        logging.info('host: {} {} begin save'.format(host_name, 'mda_statistic_data'))

    today = date.today()
    for item in mda_statistic_data:

        #存入数据库
        mda_statistic = MdaStatistic(
            city = city,
            host_name = host_name,
            host_ip = item[1],
            card_type = item[2],
            card_num = item[3],
            date_time = today
        )

        db.session.add(mda_statistic)
    db.session.commit()
    db.session.close()

def save_load_statistic(city, host_name, config):
    '''业务负荷统计表-按端口统计入库'''

    # app = create_app('production')
    # app.app_context().push()
    today_data_count = LoadStatistic.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('load_statistic host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    res = get_statistic_data(config)
    if res:
        logging.info('host: {} {} begin save'.format(host_name, 'load_statistic_data'))

    today = date.today()
    for item in res:

        #存入数据库
        load_statistic = LoadStatistic(
            city = city,
            host_name = host_name,
            host_ip = item[2],
            port = item[1],
            port_dk = item[3],
            in_utilization = item[4],
            out_utilization = item[5],
            ies_3000_user_num = item[6],
            ies_3000_utilization = item[7],
            vprn_4015_user_num = item[8],
            vprn_4015_utilization = item[9],
            date_time = today
        )

        db.session.add(load_statistic)
    db.session.commit()
    db.session.close()

def save_load_statistic_host(city, host_name, config):
    '''业务负荷统计表-按设备统计入库'''

    # app = create_app('production')
    # app.app_context().push()
    today_data_count = LoadStatisticHost.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('load_statistic_host host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    load_statistic_host_data = get_statistic_host_data(city)
    if load_statistic_host_data:
        logging.info('host: {} {} begin save'.format(host_name, 'load_statistic_host_data'))

    today = date.today()
    for item in load_statistic_host_data:

        #存入数据库
        load_statistic = LoadStatisticHost(
            city = city,
            host_name = item[0],
            host_ip = item[1],
            ies_3000_num = item[2],
            ies_3000_pool_utilization = item[3],
            vprn_4015_num = item[4],
            vprn_4015_pool_utilization = item[5],
            date_time = today
        )

        db.session.add(load_statistic)
    db.session.commit()
    db.session.close()

def save_service_statistic(city, host_name, config):
    '''业务统计入库'''

    today_data_count = ServiceStatistic.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('service_statistic host: {} today is saved'.format(host_name))
        return

    service_statistic_host_data = get_service_statistic(config)
    if service_statistic_host_data:
        logging.info('host: {} {} begin save'.format(host_name, 'service_statistic'))

    today = date.today()
    for item in service_statistic_host_data:

        #存入数据库
        service_statistic = ServiceStatistic(
            city = city,
            host_name = host_name,
            port = item[0],
            lag = item[1],
            sap = item[2],
            interface = item[3],
            service = item[4],
            service_id = item[5],
            date_time = today
        )

        db.session.add(service_statistic)
    db.session.commit()
    db.session.close()
def save_address_collect(city, host_name, config):
    '''地址采集数据入库'''

    # app = create_app('production')
    # app.app_context().push()
    save_path = os.path.join('app','static', 'address_collect', '{}-{}-地址采集.xlsx'.format(host_name, date.today().strftime('%Y%m%d')))
    if os.path.exists(save_path):
        logging.info('address_collect host: {} today is saved'.format(host_name))
        return

    address_data = get_address_data(config)
    if address_data:
        logging.info('host: {} {} begin save'.format(host_name, 'address_collect_data'))

    # for item in res:
    #     #存入数据库
    #     address_collect = AddressCollect(
    #         host_name = host_name,
    #         host_ip = item[0],
    #         ip_type = item[1],
    #         function_type = item[2],
    #         is_use = item[3],
    #         ip = item[4].strNormal(0),
    #         gateway = item[5],
    #         mask = item[6],
    #         interface_name = item[7],
    #         sap_id = item[8],
    #         next_hop = item[9],
    #         ies_vprn_id = item[10],
    #         vpn_rd = item[11],
    #         vpn_rt = item[12],
    #         description = item[13],
    #         date_time = today
    #     )

    #     db.session.add(address_collect)
    # db.session.commit()
    # db.session.close()

    ILLEGAL_CHARACTERS_RE = re.compile(r'[\000-\010]|[\013-\014]|[\016-\037]')

    #新建excel
    save_path = os.path.join('app','static', 'address_collect', '{}-{}-地址采集.xlsx'.format(host_name, date.today().strftime('%Y%m%d')))
    excel = openpyxl.Workbook()
    sheet = excel.active
    sheet['A1'] = '设备名'
    sheet['B1'] = '设备IP'
    sheet['C1'] = 'IP类型'
    sheet['D1'] = '网络功能类型'
    sheet['E1'] = '是否已分配'
    sheet['F1'] = 'IP'
    sheet['G1'] = '网关'
    sheet['H1'] = '掩码'
    sheet['I1'] = '逻辑接口编号'
    sheet['J1'] = 'sap-ID'
    sheet['K1'] = '下一跳IP'
    sheet['L1'] = 'IES/VPRN编号'
    sheet['M1'] = 'VPN-RD'
    sheet['N1'] = 'VPN-RT'
    sheet['O1'] = '接口或用户描述'


    sheet.column_dimensions['A'].width = 40.0
    sheet.column_dimensions['B'].width = 20.0
    # sheet.column_dimensions['N'].width = 20.0
    # sheet.column_dimensions['Q'].width = 15.0
    # sheet.column_dimensions['R'].width = 15.0
    cur_row = 2

    # address_data = AddressCollect.query.filter_by(host_name = host_name, date_time = date.today()).all()
    for item in address_data:
        des = ILLEGAL_CHARACTERS_RE.sub(r'', item[13])
        sheet['A'+ str(cur_row)] = host_name
        sheet['B'+ str(cur_row)] = item[0]
        sheet['C'+ str(cur_row)] = item[1]
        sheet['D'+ str(cur_row)] = item[2]
        sheet['E'+ str(cur_row)] = item[3]
        sheet['F'+ str(cur_row)] = item[4].strNormal(0)
        sheet['G'+ str(cur_row)] = item[5]
        sheet['H'+ str(cur_row)] = item[6]
        sheet['I'+ str(cur_row)] = item[7]
        sheet['J'+ str(cur_row)] = item[8]
        sheet['K'+ str(cur_row)] = item[9]
        sheet['L'+ str(cur_row)] = item[10]
        sheet['M'+ str(cur_row)] = item[11]
        sheet['N'+ str(cur_row)] = item[12]
        sheet['O'+ str(cur_row)] = des

        cur_row += 1
    
    excel.save(save_path)

def save_zuxun(city, host_name, config):
    '''保存组巡数据'''

    # app = create_app('production')
    # app.app_context().push()
    today_data_count = ZuXun.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('zuxun host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    res = zuxun_check(config)
    if res:
        logging.info('host: {} {} begin save'.format(host_name, 'zuxun'))

    today = date.today()
    zuxun = ZuXun(
        city = city,
        host_name = host_name,
        err = res,
        date_time = today
    )

    db.session.add(zuxun)
    db.session.commit()
    db.session.close()

def save_netflow(city, host_name, config):
    '''保存重要网络流量数据'''

    # app = create_app('production')
    # app.app_context().push()
    today_data_count = NetFlow.query.filter_by(site_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    res = get_netflow(config)
    if res:
        logging.info('host: {} {} begin save'.format(host_name, 'netflow'))

    today = date.today()
    netflow = NetFlow(
        carrier = 'CMCC',
        province = 'JS',
        city = city,
        network = 'CMNET',
        site_name = host_name,
        port_num_10g = res[0],
        port_utilization_10g = res[1],
        port_num_100g = res[2],
        port_utilization_100g = res[3],
        pea_uplink_throughput_utilization = res[4],
        date_time = today
    )

    db.session.add(netflow)
    db.session.commit()
    db.session.close()

def save_nat_port(city, host_name, config):
    '''保存nat port 利用率'''

    today_data_count = NatPort.query.filter_by(host_name = host_name, date_time = date.today()).count()

    if today_data_count:
        logging.info('nat port host: {} today is saved'.format(host_name))
        # db.session.close()
        return

    res = get_nat_port(config)
    if res:
        logging.info('nat port host: {} begin save'.format(host_name))

    today = date.today()
    for i in res:
        nat_port = NatPort(
            city = city,
            host_name = host_name,
            port = i[0],
            utilization = i[1],
            date_time = today
        )

        db.session.add(nat_port)
        db.session.commit()
    db.session.close()

def save_install_base():
    '''保存installbase数据'''

    # app = create_app('production')
    # app.app_context().push()

    today_data_count = InstallBase.query.filter_by(date_time = date.today()).count()
    if today_data_count:
        logging.info('install_base today is saved')
        # db.session.close()
        return

    data = dict()
    city_list = get_city_list()
    for city in city_list:
        # logging.info('city: {}'.format(city))
        for host in get_host(city):
            # logging.info('host: {}'.format(host))
            today_log = get_log_first(city, host)
            # host_name = get_host_name(today_log)
            if today_log:
                host_type, host_model, version, note = get_install_base(today_log)
                key = host_type + ' ' + host_model + ' ' + version + ' ' + note
                if not key.strip():
                    logging.info('{} get installbase err'.format(host))
                    continue
                if key in data:
                    data[key] += 1
                else:
                    data[key] = 1
            else:
                logging.error('install_base host: {} log not found'.format(host))

    logging.info('installbase begin save')
    today = date.today()
    for k, v in data.items():
        keys = k.split(' ')
        if len(keys) != 5:
            logging.error('installbase err: key: {}'.format(k))
            continue
        host_type = keys[0] + ' ' + keys[1]

        #存入数据库
        install_base = InstallBase(
            province = '江苏',
            operator = '移动',
            busines_type = '城域网',
            net_type = 'BRAS',
            host_type = keys[0] + ' ' + keys[1],
            host_model = keys[2],
            version = keys[3],
            number = v,
            note = keys[4],
            date_time = today
        )

        db.session.add(install_base)

    db.session.commit()
    db.session.close()

@main.route('/db_test')
def db_test():
    '''数据库测试'''
    address_collect = AddressCollect(
        host_name = 'aaaa',
        host_ip = 'sdsd',
        ip_type = 'sd',
        function_type = 'sd',
        is_use = 's',
        ip = 'sss',
        gateway = 'sss',
        mask = 'sds',
        interface_name = 'sds',
        sap_id = 'sdsd',
        next_hop = 'sdwd',
        ies_vprn_id = 'sdsd',
        vpn_rd = 'sdsds',
        vpn_rt = 'sdsxcc',
        description = 'sdasd'
    )


    db.session.add(address_collect)
    db.session.commit()
    db.session.close()


    a = AddressCollect.query.first()

    return a.host_name

