# encoding: utf-8
import docx
from docx.shared import Pt
from docx.shared import RGBColor
import os
import time
import datetime
import json
def set_font(font_obj, font, size):
    font_obj.name = font
    font_obj.size = Pt(size)

def set_color(paragraph, color):
    font = paragraph.runs[0].font
    font.color.rgb = RGBColor(color[0], color[1], color[2])

def add_line(doc, content, font_name, size):
    paragraph = doc.add_paragraph(content)
    font = paragraph.runs[0].font
    set_font(font, font_name, size)
# def make_report_tel(gz_path):
#     gz_file = os.path.join(gz_path, 'gz.json')
#     city = os.path.basename(gz_path)
#     if not os.path.exists(gz_file):
#         print(gz_file + ' ,文件不存在')
#         return False
#     doc = docx.Document('xunjian.docx')
#     area = city if city in ['CE', '教育网']  else city + '电信' 
#     doc.add_paragraph(area + '巡检报告', style='report-head')
#     doc.add_paragraph('上海贝尔7750设备巡检报告', style='report-head')
#     doc.add_paragraph()
#     doc.add_paragraph()
#     today_obj = datetime.datetime.now()
#     today = '%d年%d月%d日' % (today_obj.year, today_obj.month, today_obj.day)
#     doc.add_paragraph(today, style='report-date')
#     doc.add_page_break()
#     doc.add_heading('巡检情况汇总', 4)
#     f = open(gz_file,'r',encoding='utf-8')
#     gz_all = f.read()
#     gz_list = json.loads(gz_all)
#     for line in gz_list:
#         p_name = doc.add_paragraph(line[0], style='report-info')

#         if line[1]:
#             p_info = doc.add_paragraph(line[1], style='report-info')
#         if line[2]:
#             p_warn = doc.add_paragraph('注释：' + line[2], style='report-normal')

#         if not line[1] and not line[2]:
#             print('not a line 1, 2')
#         doc.add_paragraph()
        
#         set_color(p_name, (0, 0, 255))
#         set_color(p_info, (0, 0, 255))
#         set_color(p_warn, (255, 0, 0))

#     doc.add_heading('总结', 4)
#     doc.add_paragraph('1，为了保障%s电信城域网7750设备正常运行，请定期清理过滤网。' % city,
#         style='report-normal')
#     if 'Temperature' in gz_all:
#         doc.add_paragraph('2，板卡温度高建议清洗防尘网。',
#             style='report-normal')
#     f.close()
#     doc.save('%s电信巡检报告-%s.docx' % (city, today))

def make_report_tel2(gz, city):
    doc = docx.Document('xunjian.docx')
    area = city if city in ['CE', '教育网']  else city + '电信' 
    doc.add_paragraph(area + '巡检报告', style='report-head')
    doc.add_paragraph('上海贝尔7750设备巡检报告', style='report-head')
    doc.add_paragraph()
    doc.add_paragraph()
    today_obj = datetime.datetime.now()
    today = '%d年%d月%d日' % (today_obj.year, today_obj.month, today_obj.day)
    doc.add_paragraph(today, style='report-date')
    doc.add_page_break()
    doc.add_heading('巡检情况汇总', 4)
    for line in gz:
        p_name = doc.add_paragraph(line[0], style='report-info')

        if line[1]:
            p_info = doc.add_paragraph(line[1], style='report-info')
        if line[2]:
            p_warn = doc.add_paragraph(line[2], style='report-normal')

        if not line[1] and not line[2]:
            print('not line 1, 2')
        doc.add_paragraph()
        
        set_color(p_name, (0, 0, 255))
        set_color(p_info, (0, 0, 255))
        set_color(p_warn, (255, 0, 0))

    doc.add_heading('总结', 4)
    doc.add_paragraph('1，为了保障%s电信城域网7750设备正常运行，请定期清理过滤网。' % city,
        style='report-normal')

    for item in gz:
        if '板卡温度过高' in item[2]:
            doc.add_paragraph('2，板卡温度高建议清洗防尘网。',
                style='report-normal')
            break
    report_name = '%s电信巡检报告-%s.docx' % (city, today)
    doc.save(report_name)

    return report_name

def make_report_mob(gz_list, city):
    
    doc = docx.Document('xunjian.docx')
    area = city + '移动' 
    doc.add_paragraph(area + '巡检报告', style='report-head')
    doc.add_paragraph('上海贝尔7750设备巡检报告', style='report-head')
    doc.add_paragraph()
    doc.add_paragraph()
    today_obj = datetime.datetime.now()
    today = '%d年%d月%d日' % (today_obj.year, today_obj.month, today_obj.day)
    doc.add_paragraph(today, style='report-date')
    doc.add_page_break()
    doc.add_heading('巡检情况汇总', 4)
    
    for line in gz_list:
        p_name = doc.add_paragraph(line[0], style='report-info')

        if line[1]:
            p_info = doc.add_paragraph(line[1], style='report-info')
        if line[2]:
            p_warn = doc.add_paragraph('注：' + line[2], style='report-normal')

        if not line[1] and not line[2]:
            print('not a line 1, 2')
        doc.add_paragraph()
        
        set_color(p_name, (0, 0, 0))
        set_color(p_info, (0, 0, 255))
        set_color(p_warn, (255, 0, 255))

    doc.add_heading('总结', 4)
    doc.add_paragraph('1，为了保障%s移动城域网7750设备正常运行，请定期清理过滤网。' % city,
        style='report-normal')

    for item in gz_list:
        if 'Temperature' in item:
            doc.add_paragraph('2，板卡温度高建议清洗防尘网。',
            style='report-normal')
        break
    
    report_name = '%s移动巡检报告-%s.docx' % (city, today)
    doc.save(report_name)

    return report_name

# date_path = time.strftime('%Y-%m-%d', time.localtime(time.time()))
# while True:
#     folder = input('请输入日期(格式：年-月-日): ')
#     date_path = os.path.join(os.getcwd(),folder)
#     if not os.path.exists(date_path):
#         print('文件夹：%s ,不存在' % path)
#     else:
#         break

# city_list = os.listdir(date_path)

# for city in city_list:
#     city_path = os.path.join(date_path,city)
#     if  not os.path.isfile(city_path):
#         pass
        # make_report_tel(city_path)


# doc = docx.Document()

# doc.add_heading('7750设备故障汇总', 0)